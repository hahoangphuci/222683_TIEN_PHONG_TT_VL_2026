from __future__ import annotations

from typing import Iterator, Union

import docx
from docx.document import Document as _Document
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
from docx.text.run import Run


def _iter_block_items(parent: Union[_Document, _Cell]):
    """Yield Paragraph and Table items in document order.

    Based on the underlying XML order (w:p, w:tbl).
    """
    parent_elm = parent.element.body if isinstance(parent, _Document) else parent._tc

    for child in parent_elm.iterchildren():
        if child.tag.endswith("}p"):
            yield Paragraph(child, parent)
        elif child.tag.endswith("}tbl"):
            yield Table(child, parent)


def iter_docx_text_runs(doc: _Document) -> Iterator[Run]:
    """Iterate runs in body (paragraphs + tables) in order.

    Requirement: do NOT delete paragraphs; do NOT rebuild the document.
    """

    for block in _iter_block_items(doc):
        if isinstance(block, Paragraph):
            for run in block.runs:
                yield run
        else:
            # Table
            for row in block.rows:
                for cell in row.cells:
                    for inner in _iter_block_items(cell):
                        if isinstance(inner, Paragraph):
                            for run in inner.runs:
                                yield run
                        else:
                            # Nested tables are rare but supported.
                            for nrow in inner.rows:
                                for ncell in nrow.cells:
                                    for ninner in _iter_block_items(ncell):
                                        if isinstance(ninner, Paragraph):
                                            for run in ninner.runs:
                                                yield run


def iter_docx_paragraphs(doc: _Document) -> Iterator[Paragraph]:
    """Iterate paragraphs in body (including inside tables) in document order."""

    for block in _iter_block_items(doc):
        if isinstance(block, Paragraph):
            yield block
        else:
            for row in block.rows:
                for cell in row.cells:
                    for inner in _iter_block_items(cell):
                        if isinstance(inner, Paragraph):
                            yield inner
                        else:
                            for nrow in inner.rows:
                                for ncell in nrow.cells:
                                    for ninner in _iter_block_items(ncell):
                                        if isinstance(ninner, Paragraph):
                                            yield ninner


def iter_docx_header_footer_runs(doc: _Document) -> Iterator[Run]:
    """Iterate runs in all section headers/footers."""
    for section in doc.sections:
        for hf in (section.header, section.footer):
            # paragraphs
            for p in hf.paragraphs:
                for run in p.runs:
                    yield run
            # tables
            for tbl in hf.tables:
                for row in tbl.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            for run in p.runs:
                                yield run


def iter_docx_header_footer_paragraphs(doc: _Document) -> Iterator[Paragraph]:
    """Iterate paragraphs in all section headers/footers (including inside tables)."""
    for section in doc.sections:
        for hf in (section.header, section.footer):
            for p in hf.paragraphs:
                yield p
            for tbl in hf.tables:
                for row in tbl.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            yield p
