from __future__ import annotations

import os
import re
from typing import Callable, Optional

import docx

from ..extractor.docx_runs import (
    iter_docx_header_footer_paragraphs,
    iter_docx_header_footer_runs,
    iter_docx_paragraphs,
    iter_docx_text_runs,
)


class DocxRenderer:
    """In-place DOCX translation by replacing run.text only.

    Guarantees:
      - Does NOT delete paragraphs.
      - Does NOT rebuild the document.
      - Preserves run-level formatting (font, size, bold/italic, etc.) because we only change text.
      - Table structure and header/footer are preserved.
    """

    def __init__(
        self,
        translate_fn: Callable[[str], str],
        translate_batch_fn: Optional[Callable[[list[str]], list[str]]] = None,
    ):
        self.translate_fn = translate_fn
        self.translate_batch_fn = translate_batch_fn

    @staticmethod
    def _split_edge_whitespace(text: str) -> tuple[str, str, str]:
        """Split run text into leading whitespace, core text, trailing whitespace."""
        if text is None:
            return "", "", ""
        s = str(text)
        if not s:
            return "", "", ""
        left = len(s) - len(s.lstrip())
        right = len(s) - len(s.rstrip())
        prefix = s[:left] if left else ""
        suffix = s[len(s) - right:] if right else ""
        core = s[left:len(s) - right] if right else s[left:]
        return prefix, core, suffix

    @staticmethod
    def _run_font_name(run) -> str:
        """Best-effort font name lookup for a run.

        python-docx may return None for run.font.name when the font comes from
        run XML (w:rFonts) or style inheritance.
        """
        try:
            direct = run.font.name
            if direct:
                return str(direct)
        except Exception:
            pass

        try:
            from docx.oxml.ns import qn

            r_pr = run._element.find(qn("w:rPr"))
            if r_pr is None:
                return ""
            r_fonts = r_pr.find(qn("w:rFonts"))
            if r_fonts is None:
                return ""
            for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
                val = r_fonts.get(qn(attr))
                if val:
                    return str(val)
        except Exception:
            pass
        return ""

    @staticmethod
    def _is_symbol_font_name(font_name: str) -> bool:
        name = str(font_name or "").strip().lower()
        if not name:
            return False
        return any(k in name for k in (
            "symbol",
            "wingdings",
            "webdings",
            "zapfdingbats",
            "dingbats",
            "mt extra",
        ))

    def _translate_text_batch(self, texts: list[str]) -> list[str]:
        """Translate one batch of runs, with safe fallback to per-run calls."""
        if not texts:
            return []

        if self.translate_batch_fn is not None:
            try:
                out = self.translate_batch_fn(texts)
                if isinstance(out, list) and len(out) == len(texts):
                    return [str(x) if x is not None else "" for x in out]
            except Exception:
                # Fallback below keeps pipeline running even if batch parsing fails.
                pass

        return [self.translate_fn(t) for t in texts]

    def translate_docx(
        self,
        input_path: str,
        output_path: str,
        *,
        progress_cb: Optional[Callable[[int, str], None]] = None,
    ) -> str:
        doc = docx.Document(input_path)

        # Batch runs to reduce OpenRouter calls while preserving run-level formatting.
        try:
            batch_size = int(os.getenv("DOCX_RUN_BATCH_SIZE", "30"))
        except Exception:
            batch_size = 30
        batch_size = max(1, min(200, batch_size))

        entries: list[tuple[object, str, str, str]] = []

        def _collect(run):
            raw = run.text
            if raw is None:
                return
            prefix, core, suffix = self._split_edge_whitespace(str(raw))
            if not core or not core.strip():
                return
            # Keep internal newlines as-is; they carry layout/list structure in many PDFs.
            entries.append((run, prefix, core, suffix))

        for run in iter_docx_text_runs(doc):
            _collect(run)
        for run in iter_docx_header_footer_runs(doc):
            _collect(run)

        total = len(entries)
        done = 0

        for start in range(0, total, batch_size):
            chunk = entries[start:start + batch_size]
            src_batch = [item[2] for item in chunk]
            dst_batch = self._translate_text_batch(src_batch)
            if len(dst_batch) != len(chunk):
                dst_batch = [self.translate_fn(t) for t in src_batch]

            for (run, prefix, _core, suffix), dst in zip(chunk, dst_batch):
                translated = str(dst) if dst is not None else ""
                if not translated.strip():
                    translated = str(_core)
                run.text = f"{prefix}{translated}{suffix}"

                # If source run uses a symbol font but translated text contains
                # normal letters, keep size/style but switch to a Unicode text font
                # to avoid rendering tofu squares in output PDF.
                src_font_name = self._run_font_name(run)
                if self._is_symbol_font_name(src_font_name) and re.search(r"[A-Za-zÀ-ỹ]", translated):
                    try:
                        run.font.name = "Times New Roman"
                    except Exception:
                        pass

                done += 1

            if progress_cb and total > 0:
                pct = 5 + int((done / total) * 93)
                progress_cb(min(98, pct), f"DOCX: translated {done}/{total} runs (batch={batch_size})")

        doc.save(output_path)
        if progress_cb:
            progress_cb(100, f"DOCX: completed (batch={batch_size})")
        return output_path

    def translate_docx_bilingual_newline_paragraph(
        self,
        input_path: str,
        output_path: str,
        *,
        progress_cb: Optional[Callable[[int, str], None]] = None,
    ) -> str:
        """Bilingual mode (newline) rendered per paragraph.

        Why: DOCX content often splits one sentence into many runs; doing "src\ntrg"
        per-run creates many unexpected line breaks.
        Strategy:
          - Keep source paragraph runs unchanged.
          - Translate the paragraph text once.
          - Append "\n<translation>" to the last non-empty run in that paragraph.
        """

        doc = docx.Document(input_path)
        done = 0

        def _process_paragraph(p):
            nonlocal done
            runs = list(p.runs)
            if not runs:
                return
            src_para = "".join((r.text or "") for r in runs)
            if not src_para.strip():
                return
            dst_para = self.translate_fn(src_para)
            if not str(dst_para).strip():
                return

            # Find last run that has some text to attach translation.
            last = None
            for r in reversed(runs):
                if (r.text or "").strip():
                    last = r
                    break
            if last is None:
                last = runs[-1]

            last.text = f"{last.text or ''}\n{dst_para}"
            done += 1
            if progress_cb and done % 30 == 0:
                progress_cb(min(98, 5 + done // 10), f"DOCX: bilingual-newline {done} paragraphs")

        for p in iter_docx_paragraphs(doc):
            _process_paragraph(p)

        for p in iter_docx_header_footer_paragraphs(doc):
            _process_paragraph(p)

        doc.save(output_path)
        if progress_cb:
            progress_cb(100, "DOCX: completed (bilingual newline by paragraph)")
        return output_path

    def translate_docx_bilingual_sentence_inline_paragraph(
        self,
        input_path: str,
        output_path: str,
        *,
        delimiter: str = "|",
        progress_cb: Optional[Callable[[int, str], None]] = None,
        sentence_splitter: Optional[Callable[[str], list[tuple[str, str]]]] = None,
        should_pair: Optional[Callable[[str], bool]] = None,
    ) -> str:
        """Bilingual mode (sentence inline) rendered per paragraph.

        Output format per sentence: "<src> <delimiter> <dst>".
        Keeps paragraph boundaries (no merging paragraphs).
        """

        doc = docx.Document(input_path)
        done = 0

        def _default_splitter(text: str) -> list[tuple[str, str]]:
            # Fallback: treat whole paragraph as one unit.
            return [(text, "")]

        splitter = sentence_splitter or _default_splitter
        pair_decider = should_pair or (lambda _s: True)
        d = (delimiter or "|").strip() or "|"
        if len(d) > 10:
            d = d[:10]

        def _process_paragraph(p):
            nonlocal done
            runs = list(p.runs)
            if not runs:
                return
            src_para = "".join((r.text or "") for r in runs)
            if not src_para.strip():
                return

            parts = splitter(src_para)
            out_chunks: list[str] = []
            for sentence, sep in parts:
                s = sentence or ""
                if not s.strip():
                    out_chunks.append(s + (sep or ""))
                    continue
                core = s.strip()
                dst = self.translate_fn(core)
                if pair_decider(core):
                    out_chunks.append(f"{core} {d} {str(dst).strip()}{sep or ''}")
                else:
                    out_chunks.append(f"{str(dst).strip()}{sep or ''}")

            out_text = "".join(out_chunks)
            if not out_text.strip():
                return

            # Preserve formatting: write into first content run if uniform,
            # otherwise detect different formatting and write to first run only
            # (bilingual inline mode concatenates everything, format-group
            # strategy doesn't apply here; just use first-content-run).
            from docx.oxml.ns import qn as _qn
            content_indices = [i for i, r in enumerate(runs) if (r.text or "").strip()]
            if not content_indices:
                content_indices = [0]

            target = content_indices[0]
            runs[target].text = out_text
            for i, r in enumerate(runs):
                if i != target:
                    r.text = ""

            done += 1
            if progress_cb and done % 25 == 0:
                progress_cb(min(98, 5 + done // 8), f"DOCX: bilingual-sentence {done} paragraphs")

        for p in iter_docx_paragraphs(doc):
            _process_paragraph(p)

        for p in iter_docx_header_footer_paragraphs(doc):
            _process_paragraph(p)

        doc.save(output_path)
        if progress_cb:
            progress_cb(100, "DOCX: completed (bilingual sentence inline by paragraph)")
        return output_path

    def translate_docx_bilingual_inline_paragraph(
        self,
        input_path: str,
        output_path: str,
        *,
        delimiter: str = "|",
        progress_cb: Optional[Callable[[int, str], None]] = None,
        split_prefix: Optional[Callable[[str], tuple[str, str]]] = None,
    ) -> str:
        """Bilingual mode (inline) rendered per paragraph with a SINGLE delimiter.

        Output format per paragraph: "<src_para> <delimiter> <dst_para>".
        This avoids multiple separators caused by run splitting.
        """

        doc = docx.Document(input_path)
        done = 0

        d = (delimiter or "|").strip() or "|"
        if len(d) > 10:
            d = d[:10]

        def _process_paragraph(p):
            nonlocal done
            runs = list(p.runs)
            if not runs:
                return

            src_para = "".join((r.text or "") for r in runs)
            if not src_para.strip():
                return

            # Avoid double-appending if the paragraph already contains a bilingual delimiter.
            if f" {d} " in src_para:
                return

            prefix = ""
            body = src_para
            if split_prefix is not None:
                try:
                    prefix, body = split_prefix(src_para)
                except Exception:
                    prefix, body = "", src_para

            dst_para = self.translate_fn(str(body))
            dst_core = (str(dst_para) if dst_para is not None else "").strip()
            if not dst_core:
                return

            # Append to the last non-empty run to preserve existing run formatting,
            # tab stops / leaders, and other Word layout behaviors.
            last = None
            for r in reversed(runs):
                if (r.text or "").strip():
                    last = r
                    break
            if last is None:
                last = runs[-1]

            spacer = "" if (last.text or "").endswith((" ", "\t")) else " "
            last.text = f"{last.text or ''}{spacer}{d} {dst_core}"

            done += 1
            if progress_cb and done % 25 == 0:
                progress_cb(min(98, 5 + done // 8), f"DOCX: bilingual-inline {done} paragraphs")

        for p in iter_docx_paragraphs(doc):
            _process_paragraph(p)

        for p in iter_docx_header_footer_paragraphs(doc):
            _process_paragraph(p)

        doc.save(output_path)
        if progress_cb:
            progress_cb(100, "DOCX: completed (bilingual inline by paragraph)")
        return output_path

    def translate_docx_bilingual_inline_per_line(
        self,
        input_path: str,
        output_path: str,
        *,
        delimiter: str = "|",
        progress_cb: Optional[Callable[[int, str], None]] = None,
    ) -> str:
        """Bilingual inline mode optimised for pdf2docx-generated DOCX files.

        pdf2docx merges multiple logical PDF lines into a single DOCX
        paragraph connected by runs whose text is exactly ``\\n``.  This
        method splits each paragraph at those newline-runs, translates
        each *logical line* independently, and appends the translation
        right after that line so positions stay correct.

        For paragraphs that have no embedded ``\\n`` runs the behaviour is
        the same as ``translate_docx_bilingual_inline_paragraph``.
        """
        doc = docx.Document(input_path)
        done = 0

        d = (delimiter or "|").strip() or "|"
        if len(d) > 10:
            d = d[:10]

        # ── Symbol-font Private-Use-Area → standard Unicode ──
        _SYMBOL_MAP = {
            "\uf02b": "+",   # SymbolMT "+"
            "\uf0b7": "\u2022",  # bullet •
            "\uf0a7": "\u00a7",  # section §
        }

        import re
        _heading_start_re = re.compile(r'^[+\-](\s|$)')

        # ── Vietnamese diacritics set for OCR merging detection ──
        # Tone-marked vowels: these only appear as the NUCLEUS of a
        # Vietnamese syllable.  When one is directly followed by an
        # initial-consonant cluster + another vowel and no space in
        # between, two syllables have been stuck together by pdf2docx.
        _VN_TONE_VOWELS = (
            "àáảãạăắằẳẵặâấầẩẫậèéẻẽẹêếềểễệ"
            "ìíỉĩịòóỏõọôốồổỗộơớờởỡợùúủũụ"
            "ưứừửữựỳýỷỹỵđ"
            "ÀÁẢÃẠĂẮẰẲẴẶÂẤẦẨẪẬÈÉẺẼẸÊẾỀỂỄỆ"
            "ÌÍỈĨỊÒÓỎÕỌÔỐỒỔỖỘƠỚỜỞỠỢÙÚỦŨỤ"
            "ƯỨỪỬỮỰỲÝỶỸỴĐ"
        )
        _VN_ALL_VOWELS = (
            "aăâeêioôơuưyAĂÂEÊIOÔƠUƯY"
            + _VN_TONE_VOWELS
        )
        # Regex: diacritic vowel stuck to initial consonant of next syllable.
        #   Group 1 = tone vowel (end of prev syllable)
        #   Group 2 = initial consonant cluster (start of next syllable)
        #   Lookahead = a vowel (confirming it IS an initial, not a final)
        _STUCK_SYLLABLE_RE = re.compile(
            r'([' + re.escape(_VN_TONE_VOWELS) + r'])'
            r'(ngh|nh|ng|ch|gh|gi|kh|ph|qu|th|tr|[bcdđghklmnpqrstvx])'
            r'(?=[' + re.escape(_VN_ALL_VOWELS) + r'])',
            re.IGNORECASE,
        )
        _LETTER_OR_VN = re.compile(r'[a-zA-Z\u00C0-\u024F\u1E00-\u1EFF]')

        def _fix_ocr_text(text: str) -> str:
            """Fix common pdf2docx / OCR extraction artefacts.

            - Insert missing space between stuck Vietnamese syllables
              e.g. "Hệthống" → "Hệ thống", "cửviên" → "cử viên"
            - Fix stuck text after closing brackets: ")kết" → ") kết"
            - Fix stuck text before opening brackets: "viên(" → "viên ("
            - Fix list marker spacing: "+Sach" → "+ Sach"
            - Collapse multiple spaces
            """
            if not text or len(text) < 2:
                return text

            s = text
            # 1. Fix stuck Vietnamese syllables (main OCR issue)
            s = _STUCK_SYLLABLE_RE.sub(r'\1 \2', s)

            # 2. Closing bracket/paren stuck to next word
            s = re.sub(r'([)\]}>])(' + r'[a-zA-Z\u00C0-\u024F\u1E00-\u1EFF])', r'\1 \2', s)

            # 3. Letter stuck to opening bracket
            s = re.sub(r'([a-zA-Z\u00C0-\u024F\u1E00-\u1EFF])(\()', r'\1 \2', s)

            # 4. List marker spacing: "+Text" → "+ Text" (only at start
            #    of string or after whitespace)
            s = re.sub(r'(^|(?<=\s))([+\-])([a-zA-Z\u00C0-\u024F\u1E00-\u1EFF])', r'\2 \3', s)

            # 5. Collapse multiple spaces
            s = re.sub(r'  +', ' ', s)
            return s

        def _clean_translation(text: str) -> str:
            """Post-process AI translation output."""
            if not text:
                return text
            s = text.strip()
            # Remove leading pipe or delimiter
            if s.startswith('|'):
                s = s[1:].strip()
            # Remove wrapping quotes
            if len(s) >= 2 and s[0] == s[-1] and s[0] in '"\'':
                s = s[1:-1].strip()
            # Remove markdown formatting
            s = s.replace('**', '')
            # Collapse whitespace
            s = re.sub(r'  +', ' ', s)
            return s.strip()

        def _is_bullet_font(fn: str) -> bool:
            return ("Courier" in fn or "Symbol" in fn
                    or "Wingding" in fn)

        def _process_paragraph(p):
            nonlocal done
            runs = list(p.runs)
            if not runs:
                return

            # ── Pre-process: fix SymbolMT Private-Use-Area chars ──
            for ri, r in enumerate(runs):
                fn = r.font.name or ""
                if "Symbol" in fn:
                    old = r.text or ""
                    new = old
                    for k, v in _SYMBOL_MAP.items():
                        new = new.replace(k, v)
                    if new != old:
                        r.text = new
                        r.font.name = None  # inherit paragraph font
                        # If this run is just a "+" marker, add trailing
                        # space so it won't stick to the next run.
                        if new.strip() in ('+', '-', '\u2022'):
                            nxt_text = ""
                            if ri + 1 < len(runs):
                                nxt_text = (runs[ri + 1].text or "")
                            if nxt_text and nxt_text[0:1].isalpha():
                                r.text = new.rstrip() + " "

            # ── Pre-process: fix OCR artefacts in every run ──
            for r in runs:
                old = r.text or ""
                if old and old.strip():
                    fixed = _fix_ocr_text(old)
                    if fixed != old:
                        r.text = fixed

            # ── Detect bullet paragraph ──
            # A bullet paragraph starts with a bullet marker ("o", "•")
            # in CourierNewPSMT / SymbolMT.  Its \n and \t runs are
            # visual wraps within one bullet point, NOT logical breaks.
            _is_bullet = False
            for r in runs:
                rt = (r.text or "").strip()
                if not rt:
                    continue
                fn = r.font.name or ""
                if len(rt) <= 2 and _is_bullet_font(fn):
                    _is_bullet = True
                break  # only check first non-empty run

            if _is_bullet:
                # ── Bullet paragraph: keep as one logical block ──
                # Collect all content runs (skip bullet, tabs, newlines).
                # Split off any merged heading (e.g. "+ PhieuMuon" stuck
                # to end of content run) as a separate group.
                content_runs = []
                heading_runs = []
                _in_heading = False
                for r in runs:
                    rt = (r.text or "").strip()
                    fn = r.font.name or ""
                    # Skip whitespace-only runs (tabs, newlines)
                    if not rt:
                        continue
                    # Skip bullet marker itself
                    if len(rt) <= 2 and _is_bullet_font(fn):
                        continue
                    # Detect merged heading at end of bullet content
                    if (_heading_start_re.match(r.text or "")
                            and content_runs):
                        _in_heading = True
                        # Insert visual line break before heading
                        r.text = "\n" + (r.text or "")
                    if _in_heading:
                        heading_runs.append(r)
                    else:
                        content_runs.append(r)
                groups = []
                if content_runs:
                    groups.append(content_runs)
                if heading_runs:
                    groups.append(heading_runs)
            else:
                # ── Non-bullet: split at \n; tab → split only for
                #    heading markers (+/-), otherwise keep as wrapping ──
                groups: list[list] = [[]]
                for ri, r in enumerate(runs):
                    txt = r.text or ""
                    if "\n" in txt and not txt.replace("\n", "").strip():
                        # Pure newline separator → always split
                        groups.append([])
                    elif (txt.replace("\t", "").strip() == ""
                          and "\t" in txt
                          and groups[-1]  # has preceding content
                          and ri + 1 < len(runs)
                          and (runs[ri + 1].text or "").strip()):
                        next_txt = (runs[ri + 1].text or "").strip()
                        if _heading_start_re.match(next_txt):
                            # Next content is a section heading → split
                            r.text = "\n"
                            groups.append([])
                        else:
                            # Tab is visual line wrapping → keep in group
                            groups[-1].append(r)
                    else:
                        groups[-1].append(r)

            # Filter empty groups
            groups = [g for g in groups if g]
            if not groups:
                return

            # ── Translate each group ──
            for grp in groups:
                line_text = "".join((r.text or "") for r in grp).strip()
                if not line_text:
                    continue
                # Ignore runs that are only tabs/spaces
                core_text = line_text.replace("\t", "").strip()
                if not core_text:
                    continue
                # Skip very short fragments (bullet markers like "o")
                if len(core_text) < 3:
                    continue
                # Skip already-translated lines
                if f" {d} " in line_text:
                    continue

                # Apply OCR fix to the joined text sent to translator
                clean_text = _fix_ocr_text(line_text)
                dst = self.translate_fn(clean_text)
                dst_core = _clean_translation(
                    str(dst) if dst is not None else ""
                )
                if not dst_core:
                    continue

                # Find last content run in this group
                last = None
                for r in reversed(grp):
                    if (r.text or "").strip():
                        last = r
                        break
                if last is None:
                    last = grp[-1]

                spacer = "" if (last.text or "").endswith((" ", "\t")) else " "
                last.text = f"{last.text or ''}{spacer}{d} {dst_core}"
                done += 1

            if progress_cb and done % 15 == 0:
                progress_cb(min(98, 5 + done // 5),
                            f"DOCX: bilingual-per-line {done} lines")

        for p in iter_docx_paragraphs(doc):
            _process_paragraph(p)

        for p in iter_docx_header_footer_paragraphs(doc):
            _process_paragraph(p)

        doc.save(output_path)
        if progress_cb:
            progress_cb(100, "DOCX: completed (bilingual per-line)")
        return output_path

    def translate_docx_bilingual_sentence_newline_paragraph(
        self,
        input_path: str,
        output_path: str,
        *,
        progress_cb: Optional[Callable[[int, str], None]] = None,
        sentence_splitter: Optional[Callable[[str], list[tuple[str, str]]]] = None,
        should_pair: Optional[Callable[[str], bool]] = None,
    ) -> str:
        """Bilingual mode (sentence newline) rendered per paragraph.

        Required format per sentence:
          Line 1: source sentence
          Line 2: translation

        Notes:
          - Keeps paragraph boundaries (does not merge/split paragraphs).
          - Preserves original newlines that exist between sentences.
        """

        doc = docx.Document(input_path)
        done = 0

        def _default_splitter(text: str) -> list[tuple[str, str]]:
            return [(text, "")]

        splitter = sentence_splitter or _default_splitter
        pair_decider = should_pair or (lambda _s: True)

        def _norm_sep(sep: str) -> str:
            if sep and ("\n" in sep or "\r" in sep):
                return sep
            return "\n" if sep is not None else ""

        def _process_paragraph(p):
            nonlocal done
            runs = list(p.runs)
            if not runs:
                return
            src_para = "".join((r.text or "") for r in runs)
            if not src_para.strip():
                return

            parts = splitter(src_para)
            out_chunks: list[str] = []
            for sentence, sep in parts:
                raw = sentence or ""
                if not raw.strip():
                    # Keep blank fragments as-is
                    out_chunks.append(raw + (sep or ""))
                    continue
                core = raw.strip()
                dst = self.translate_fn(core)
                if pair_decider(core):
                    out_chunks.append(f"{core}\n{str(dst).strip()}{_norm_sep(sep or '')}")
                else:
                    out_chunks.append(f"{str(dst).strip()}{_norm_sep(sep or '')}")

            out_text = "".join(out_chunks)
            if not out_text.strip():
                return

            runs[0].text = out_text
            for r in runs[1:]:
                r.text = ""

            done += 1
            if progress_cb and done % 25 == 0:
                progress_cb(min(98, 5 + done // 8), f"DOCX: bilingual-sentence-newline {done} paragraphs")

        for p in iter_docx_paragraphs(doc):
            _process_paragraph(p)

        for p in iter_docx_header_footer_paragraphs(doc):
            _process_paragraph(p)

        doc.save(output_path)
        if progress_cb:
            progress_cb(100, "DOCX: completed (bilingual sentence newline by paragraph)")
        return output_path

    def translate_docx_bilingual_paren_paragraph(
        self,
        input_path: str,
        output_path: str,
        *,
        progress_cb: Optional[Callable[[int, str], None]] = None,
        should_pair: Optional[Callable[[str], bool]] = None,
        formatter: Optional[Callable[[str], tuple[str, str]]] = None,
    ) -> str:
        """Bilingual mode using parentheses: "<src> (<dst>)".

        Keeps paragraph boundaries and preserves run formatting by only appending
        the translation to the last non-empty run.
        """

        doc = docx.Document(input_path)
        done = 0
        pair_decider = should_pair or (lambda _s: True)

        def _process_paragraph(p):
            nonlocal done
            runs = list(p.runs)
            if not runs:
                return

            src_para = "".join((r.text or "") for r in runs)
            if not src_para.strip():
                return

            core = src_para.strip()
            action = None
            payload = None

            if formatter is not None:
                try:
                    action, payload = formatter(core)
                except Exception:
                    action, payload = None, None

            if not action:
                dst = self.translate_fn(core)
                dst_core = (str(dst) if dst is not None else "").strip()
                if not dst_core:
                    return

                if pair_decider(core):
                    action, payload = "append", dst_core
                else:
                    action, payload = "replace", dst_core

            last = None
            for r in reversed(runs):
                if (r.text or "").strip():
                    last = r
                    break
            if last is None:
                last = runs[-1]

            if action == "append":
                last.text = f"{last.text or ''} ({str(payload or '').strip()})"
            else:
                runs[0].text = str(payload or "")
                for r in runs[1:]:
                    r.text = ""

            done += 1
            if progress_cb and done % 40 == 0:
                progress_cb(min(98, 5 + done // 12), f"DOCX: paren bilingual {done} paragraphs")

        for p in iter_docx_paragraphs(doc):
            _process_paragraph(p)
        for p in iter_docx_header_footer_paragraphs(doc):
            _process_paragraph(p)

        doc.save(output_path)
        if progress_cb:
            progress_cb(100, "DOCX: completed (bilingual paren by paragraph)")
        return output_path
