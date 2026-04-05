import os
import re
import unicodedata
import time
import uuid
import io
import zipfile
import shutil
import docx
import openpyxl
from werkzeug.utils import secure_filename
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.shared import Inches, RGBColor


class ProviderRateLimitError(Exception):
    """Raised when the upstream AI provider indicates a hard rate limit (429 or insufficient credits)."""


class FileService:
    def __init__(self, translator=None, ocr_image_to_text=None, ocr_translate_overlay=None, ocr_image_to_bboxes=None):
        """translator: callable(text, source_lang, target_lang) -> translated_text

        ocr_image_to_text: optional callable(image_path, ocr_langs=None) -> text
        """
        from concurrent.futures import ThreadPoolExecutor

        self.upload_folder = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(__file__))), 'uploads')
        self.download_folder = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(__file__))), 'downloads')
        os.makedirs(self.upload_folder, exist_ok=True)
        os.makedirs(self.download_folder, exist_ok=True)
        self.translator = translator
        self.ocr_image_to_text = ocr_image_to_text
        self.ocr_translate_overlay = ocr_translate_overlay
        self.ocr_image_to_bboxes = ocr_image_to_bboxes
        self.has_tesseract = False  # Will be set by TranslationService
        # Performance tuning
        try:
            from app import config as app_config
            self.concurrency = getattr(app_config.Config, 'TRANSLATION_CONCURRENCY', 4)
            self.retries = getattr(app_config.Config, 'TRANSLATION_RETRIES', 3)
            self.backoff = getattr(app_config.Config, 'TRANSLATION_BACKOFF', 1.5)
        except Exception:
            self.concurrency = int(os.getenv('TRANSLATION_CONCURRENCY', '4'))
            self.retries = int(os.getenv('TRANSLATION_RETRIES', '3'))
            self.backoff = float(os.getenv('TRANSLATION_BACKOFF', '1.5'))
        self._executor_cls = ThreadPoolExecutor

    def _translate_with_retry(self, text, target_lang, *, context=None):
        """Translate a piece of text with retry/backoff on transient errors.

        IMPORTANT: If a provider rate-limit or "insufficient credits" error is encountered,
        fail fast by raising ProviderRateLimitError so the calling job can abort immediately
        instead of continuing and wasting quota/retries.
        """
        if not self.translator:
            raise RuntimeError('Translator not configured')
        last = None
        attempt = 0
        max_attempts = max(1, self.retries)
        while attempt < max_attempts:
            try:
                try:
                    out = self.translator(text, 'auto', target_lang, context=context)
                except TypeError:
                    # Backward compatibility for translator callbacks that only accept 3 params.
                    out = self.translator(text, 'auto', target_lang)
                return out
            except Exception as e:
                last = e
                err = str(e).lower()
                # Fail fast for provider rate limits or insufficient credits
                if any(k in err for k in ('429', 'too many requests', 'rate', 'insufficient', 'free-models', '402', 'credit', 'requires more credits')):
                    try:
                        print(f"Provider rate limit or insufficient credits encountered: {e}")
                    except UnicodeEncodeError:
                        print("Provider rate limit encountered: ", repr(e))
                    raise ProviderRateLimitError(str(e))
                # Retry on transient network errors
                if any(k in err for k in ('temporarily', 'timed out', 'timeout', 'connection')):
                    sleep_time = (self.backoff ** attempt)
                    print(f"Translate retry {attempt+1}/{self.retries} after {sleep_time}s due to: {e}")
                    time.sleep(sleep_time)
                    attempt += 1
                    continue
                # Non-retryable errors: break
                break
        # Final attempt to raise helpful error
        raise last

    def process_document(self, file_path, target_lang, progress_callback=None, *, ocr_images=False, ocr_langs=None, ocr_mode=None, bilingual_mode=None, bilingual_delimiter=None):
        filename = os.path.basename(file_path)
        name, ext = os.path.splitext(filename)

        if ext.lower() == '.pdf':
            return self._process_pdf(file_path, target_lang, progress_callback, bilingual_mode=bilingual_mode, bilingual_delimiter=bilingual_delimiter)
        elif ext.lower() == '.docx':
            return self._process_docx(file_path, target_lang, progress_callback, ocr_images=ocr_images, ocr_langs=ocr_langs, ocr_mode=ocr_mode, bilingual_mode=bilingual_mode, bilingual_delimiter=bilingual_delimiter)
        elif ext.lower() == '.xlsx':
            return self._process_xlsx(file_path, target_lang, progress_callback)
        elif ext.lower() == '.txt':
            return self._process_txt(file_path, target_lang, progress_callback)
        else:
            raise ValueError("Unsupported file type")

    def _process_pdf(self, file_path, target_lang, progress_callback=None, *, bilingual_mode=None, bilingual_delimiter=None):
        """Translate a text-based PDF while preserving original layout.

        Uses the layout-aware PdfLayoutRenderer which:
          1. Detects document layout (titles, paragraphs, tables, lists, form fields, dotted placeholders)
          2. Translates text preserving structure per element type
          3. Reconstructs layout preserving exact spacing, line breaks, alignment, table structure

        Bilingual modes:
          - None / 'none': normal (replace original with translation)
          - 'inline':  song ngữ liền kề — "Original | Translated" in same box
          - 'newline': song ngữ xuống dòng — original on top, translated below
          - 'preserve_layout': dịch song ngữ liền kề (tương đương 'inline')
          - 'line_by_line': dịch song ngữ xuống dòng (tương đương 'newline')
        """

        try:
            import fitz  # PyMuPDF
        except Exception as e:
            raise RuntimeError("PyMuPDF is required for PDF translation. Install 'PyMuPDF'.") from e

        if not os.path.exists(file_path):
            raise FileNotFoundError(file_path)

        target_key = str(target_lang or "").strip().lower()
        skip_vietnamese_source = target_key in ("vi", "vi-vn", "vietnamese", "vn", "viet")

        def _is_probably_vietnamese(s: str) -> bool:
            if not s:
                return False
            core = str(s).strip().lower()
            if not core:
                return False
            return bool(re.search(r"[ăâđêôơưáàảãạấầẩẫậắằẳẵặéèẻẽẹếềểễệíìỉĩịóòỏõọốồổỗộớờởỡợúùủũụứừửữựýỳỷỹỵ]", core))

        # ── Use layout-aware PDF renderer (V2 pipeline) ──
        try:
            bi_mode_v2 = (str(bilingual_mode).strip().lower() if bilingual_mode else 'none')
            if bi_mode_v2 == 'preserve_layout':
                bi_mode_v2 = 'inline'
            elif bi_mode_v2 == 'line_by_line':
                bi_mode_v2 = 'newline'
            if bi_mode_v2 not in ('none', 'inline', 'newline'):
                bi_mode_v2 = 'none'

            bi_delim_v2 = self._normalize_bilingual_delimiter(bilingual_delimiter) if bi_mode_v2 == 'inline' else '|'

            pdf_ctx = 'document_pdf_adjacent_inline' if bi_mode_v2 == 'inline' else 'document_pdf'

            _LIST_PREFIX_RE = re.compile(
                r"^(\s*(?:"
                r"[-+*•◦▪▫○●■□☐☑✓✔◆◇]"
                r"|\d+[.)]"
                r"|[A-Za-z][.)]"
                r"|[ivxIVX]{1,4}[.)]"
                r")\s*)"
            )
            # Preserve real function/method calls but avoid masking attribute/type fragments
            # like "bookID (PK, string)" (space + typed metadata, not a callable token).
            _FUNC_CALL_RE = re.compile(
                r"\b[A-Za-z_][A-Za-z0-9_]*\([^()\n]{0,80}\)"
                r"|\b[A-Za-z_][A-Za-z0-9_]*\s*\(\s*\)"
            )
            _CODE_IDENT_RE = re.compile(
                r"\b(?:"
                r"[A-Za-z_]*_[A-Za-z0-9_]+"          # snake_case / mixed_with_underscore
                r"|[a-z]+[A-Z][A-Za-z0-9_]*"         # lowerCamelCase
                r"|[A-Z][a-z0-9]+[A-Z]{2,}[A-Za-z0-9_]*"  # UpperCamel with acronym suffix, e.g. UserID
                r")\b"
            )
            _ABBR_RE = re.compile(
                r"\b(?:PK|FK|ID|MSSV|MSV|STT|NULL|N/A|API|SQL|URL|UUID|DOB|No\.)\b",
                re.IGNORECASE,
            )
            _SQL_TYPE_RE = re.compile(
                r"\b(?:N?VARCHAR|CHAR|TEXT|INT|INTEGER|BIGINT|SMALLINT|TINYINT"
                r"|DATE|DATETIME|TIMESTAMP|BOOLEAN|BOOL|FLOAT|DOUBLE|DECIMAL)"
                r"\s*(?:\(\s*\d+(?:\s*,\s*\d+)?\s*\))?\b",
                re.IGNORECASE,
            )
            _PUNCT_RUN_RE = re.compile(r"(\.{2,}|_{2,}|-{2,}|\+{2,})")
            _SYMBOL_GLYPH_RE = re.compile(
                r"([\uE000-\uF8FF□☐☑✓✔▪▫■●○◦◆◇◻◼◽◾→←↔↦➜➤➔➟➢➣]+)"
            )
            _MASK_TOKEN_RE = re.compile(r"^\[\[(?:CODE|ID|SYM|GLY)\d{3}\]\]$")
            _MASK_LEAK_RE = re.compile(r"(\[\[(?:CODE|ID|SYM|GLY)\d{3}\]\]|__(?:CODE|ID|SYM|GLY)\d{3}__)")

            def _split_list_prefix(text):
                s = str(text or "")
                m = _LIST_PREFIX_RE.match(s)
                if not m:
                    return "", s
                return m.group(1), s[m.end():]

            def _encode_controls(text):
                return str(text or "").replace("\r\n", "[[LB]]").replace("\n", "[[LB]]").replace("\r", "[[LB]]").replace("\t", "[[TAB]]")

            def _decode_controls(text):
                s = str(text or "")
                # Accept both current and legacy placeholders for backward compatibility.
                return (
                    s.replace("[[LB]]", "\n")
                    .replace("[[TAB]]", "\t")
                    .replace("__LB__", "\n")
                    .replace("__TAB__", "\t")
                )

            def _mask_nontranslatables(text):
                s = str(text or "")
                mapping = {}
                counter = 0

                def _next_key(kind):
                    nonlocal counter
                    counter += 1
                    return f"[[{kind}{counter:03d}]]"

                def _mask_with(pattern, kind, value):
                    def _repl(m):
                        token = m.group(0)
                        if _MASK_TOKEN_RE.fullmatch(token):
                            return token
                        key = _next_key(kind)
                        mapping[key] = token
                        return key
                    return pattern.sub(_repl, value)

                # Keep function calls and code-like identifiers verbatim.
                s = _mask_with(_FUNC_CALL_RE, "CODE", s)
                s = _mask_with(_CODE_IDENT_RE, "ID", s)
                s = _mask_with(_ABBR_RE, "ID", s)
                s = _mask_with(_SQL_TYPE_RE, "ID", s)
                # Keep repeated punctuation runs (...., ___, ---, ++).
                s = _mask_with(_PUNCT_RUN_RE, "SYM", s)
                # Keep symbol bullets/checkbox glyphs verbatim.
                s = _mask_with(_SYMBOL_GLYPH_RE, "GLY", s)
                return s, mapping

            def _unmask_nontranslatables(text, mapping):
                s = str(text or "")
                for _ in range(4):
                    before = s
                    for key in sorted(mapping.keys(), key=len, reverse=True):
                        s = s.replace(key, mapping[key])
                    if s == before:
                        break
                return s

            def _has_mask_leak(text):
                return bool(_MASK_LEAK_RE.search(str(text or "")))

            def _missing_mask_tokens(text, mapping):
                if not mapping:
                    return False
                s = str(text or "")
                for key in mapping.keys():
                    if key not in s:
                        return True
                return False

            def _should_translate_body(text):
                s = str(text or "")
                if not s.strip():
                    return False
                return bool(re.search(r"[A-Za-zÀ-ỹ]", s))

            def _translate_preserving_markup(src_text, *, context_name):
                src = str(src_text or "")
                if not src:
                    return src
                prefix, body = _split_list_prefix(src)
                if not _should_translate_body(body):
                    return src

                masked_body, mapping = _mask_nontranslatables(_encode_controls(body))
                translated = self._translate_with_retry(masked_body, target_lang, context=context_name)
                translated_text = str(translated or "").strip()
                if _missing_mask_tokens(translated_text, mapping):
                    return src
                restored = _decode_controls(_unmask_nontranslatables(translated_text, mapping))
                if _has_mask_leak(restored):
                    return src
                if not restored.strip():
                    return src
                return f"{prefix}{restored}" if prefix else restored

            def _layout_translate_fn(text):
                return _translate_preserving_markup(text, context_name=pdf_ctx)

            def _layout_translate_batch_fn(lines):
                """Translate a batch of run texts in one API call (best effort)."""
                if not lines:
                    return []

                prepared = {}
                tagged_lines = []

                for i, src in enumerate(lines, start=1):
                    original = str(src or "")
                    prefix, body = _split_list_prefix(original)
                    if not _should_translate_body(body):
                        prepared[i] = {
                            'needs_translate': False,
                            'result': original,
                        }
                        continue

                    masked_body, mapping = _mask_nontranslatables(_encode_controls(body))
                    prepared[i] = {
                        'needs_translate': True,
                        'prefix': prefix,
                        'mapping': mapping,
                        'source': original,
                    }
                    tagged_lines.append(f"__R{i:03d}__ {masked_body}")

                if not tagged_lines:
                    return [str(x or "") for x in lines]

                payload = "\n".join(tagged_lines)

                try:
                    raw = self._translate_with_retry(
                        payload,
                        target_lang,
                        context='document_pdf_run_batch',
                    )
                    parsed: dict[int, str] = {}
                    pattern = re.compile(r"(?ms)^\s*__R(\d{3})__\s*(.*?)(?=^\s*__R\d{3}__|\Z)")
                    for m in pattern.finditer(str(raw or '')):
                        idx = int(m.group(1))
                        parsed[idx] = m.group(2).strip()

                    results = []
                    for i, src in enumerate(lines, start=1):
                        info = prepared.get(i)
                        if not info or not info.get('needs_translate'):
                            results.append(str(src or ""))
                            continue

                        translated_row = parsed.get(i)
                        if translated_row is None:
                            results.append(_translate_preserving_markup(str(src or ""), context_name=pdf_ctx))
                            continue

                        if _missing_mask_tokens(translated_row, info['mapping']):
                            results.append(info['source'])
                            continue

                        restored = _decode_controls(
                            _unmask_nontranslatables(translated_row, info['mapping'])
                        )
                        restored = restored.strip()
                        if _has_mask_leak(restored):
                            results.append(info['source'])
                            continue
                        if not restored:
                            results.append(info['source'])
                            continue
                        prefix = info['prefix']
                        results.append(f"{prefix}{restored}" if prefix else restored)

                    return results
                except Exception as batch_err:
                    print(f"Batch run translation fallback to per-run: {batch_err}")

                # Fallback: preserve correctness if batch output is malformed.
                return [_translate_preserving_markup(str(line or ""), context_name=pdf_ctx) for line in lines]

            from app.services.document_v2.renderer.pdf import PdfLayoutRenderer

            renderer = PdfLayoutRenderer(
                translate_fn=_layout_translate_fn,
                translate_batch_fn=_layout_translate_batch_fn,
                download_folder=self.download_folder,
                skip_vietnamese_source=skip_vietnamese_source,
                target_is_vietnamese=skip_vietnamese_source,
            )

            _skip_docx = str(os.getenv('PDF_TRANSLATE_VIA_DOCX', '1')).strip().lower() in ('0', 'false', 'no')
            _allow_inline_docx_fallback = str(os.getenv('PDF_INLINE_ALLOW_DOCX_FALLBACK', '0')).strip().lower() in ('1', 'true', 'yes', 'on')

            # ── Dedicated adjacent bilingual PDF mode (inline) ─────────────────────
            # IMPORTANT: keep normal mode ('none') untouched.
            # Use block renderer inline mode (single-line bbox drawing) to avoid
            # overlap artifacts on dense pages while keeping delimiter customisation.
            if bi_mode_v2 == 'inline':
                try:
                    from app.services.document_v2.renderer.pdf_blocks import PdfBlockTranslator
                    inline_renderer = PdfBlockTranslator(
                        translate_fn=_layout_translate_fn,
                        download_folder=self.download_folder,
                        merge_paragraphs=False,
                        bilingual_mode='inline',
                        bilingual_delimiter=bi_delim_v2,
                        inline_table_mode='translate-only',
                    )
                    out_path = inline_renderer.translate_pdf(
                        input_path=file_path,
                        progress_cb=progress_callback,
                    )
                    if progress_callback:
                        progress_callback(100, "PDF: completed (adjacent bilingual)")
                    return out_path
                except Exception as inline_err:
                    print(f"Inline PDF block renderer failed: {inline_err}")
                    if _allow_inline_docx_fallback and not _skip_docx:
                        try:
                            out_path = renderer.translate_pdf_via_docx(
                                input_path=file_path,
                                bilingual_mode='inline',
                                bilingual_delimiter=bi_delim_v2,
                                progress_cb=progress_callback,
                            )
                            if progress_callback:
                                progress_callback(100, "PDF: completed (adjacent bilingual via DOCX fallback)")
                            return out_path
                        except Exception as inline_docx_err:
                            print(f"Inline PDF DOCX fallback failed: {inline_docx_err}")
                    else:
                        print("Inline PDF DOCX fallback is disabled (set PDF_INLINE_ALLOW_DOCX_FALLBACK=1 to enable).")
                    raise

            # ── Strategy 0 (primary for non-bilingual): block-by-block pipeline ─────
            # PDF → extract blocks (text+bbox+font) → translate → render by bbox.
            # Set PDF_BLOCK_PIPELINE=0 to skip this strategy.
            _use_block_pipeline = str(os.getenv('PDF_BLOCK_PIPELINE', '1')).strip().lower() not in ('0', 'false', 'no')
            if _use_block_pipeline and bi_mode_v2 == 'none':
                try:
                    from app.services.document_v2.renderer.pdf_blocks import PdfBlockTranslator
                    block_renderer = PdfBlockTranslator(
                        translate_fn=_layout_translate_fn,
                        download_folder=self.download_folder,
                    )
                    out_path = block_renderer.translate_pdf(
                        input_path=file_path,
                        progress_cb=progress_callback,
                    )
                    if progress_callback:
                        progress_callback(100, "PDF: completed")
                    return out_path
                except Exception as block_err:
                    print(f"Block pipeline failed, falling back: {block_err}")

            # ── Strategy 1 (opt-in for non-bilingual): PDF → DOCX → runs(batch) → PDF ──
            # Disabled by default because it can alter font metrics/table geometry.
            # Set PDF_DOCX_LINE_PIPELINE=1 to enable this strategy.
            _use_docx_line_pipeline = str(os.getenv('PDF_DOCX_LINE_PIPELINE', '0')).strip().lower() not in ('0', 'false', 'no')
            if not _skip_docx and _use_docx_line_pipeline and bi_mode_v2 == 'none':
                try:
                    out_path = renderer.translate_pdf_via_docx(
                        input_path=file_path,
                        bilingual_mode='none',
                        bilingual_delimiter=bi_delim_v2,
                        progress_cb=progress_callback,
                    )
                    if progress_callback:
                        progress_callback(100, "PDF: completed (via DOCX run-batch)")
                    return out_path
                except Exception as docx_line_err:
                    print(f"PDF-via-DOCX run-batch pipeline failed, falling back: {docx_line_err}")

            # ── Strategy 2 (bilingual modes): PDF → DOCX → Translate → PDF ──────────
            # Preserves layout better because DOCX has proper paragraph/table
            # structure.  Falls back to ReportLab if conversion fails.
            # Set PDF_TRANSLATE_VIA_DOCX=0 to force ReportLab instead.
            if not _skip_docx and bi_mode_v2 in ('inline', 'newline'):
                try:
                    out_path = renderer.translate_pdf_via_docx(
                        input_path=file_path,
                        bilingual_mode=bi_mode_v2,
                        bilingual_delimiter=bi_delim_v2,
                        progress_cb=progress_callback,
                    )
                    if progress_callback:
                        progress_callback(100, "PDF: completed (via DOCX)")
                    return out_path
                except Exception as docx_err:
                    print(f"PDF-via-DOCX pipeline failed, falling back to ReportLab: {docx_err}")

            # ── Strategy 3 (fallback / non-bilingual): Direct ReportLab rendering ──
            out_path = renderer.translate_pdf(
                input_path=file_path,
                bilingual_mode=bi_mode_v2,
                bilingual_delimiter=bi_delim_v2,
                progress_cb=progress_callback,
            )
            if progress_callback:
                progress_callback(100, "PDF: completed")
            return out_path
        except Exception as layout_err:
            # Fallback to legacy pipeline if layout renderer fails
            print(f"Layout-aware PDF renderer failed, falling back to legacy: {layout_err}")

        # ── Legacy fallback: direct redact+insert pipeline ──

        # ── Bilingual mode ──
        bi_mode = (str(bilingual_mode).strip().lower() if bilingual_mode else 'none')
        # Normalize new mode names
        if bi_mode == 'preserve_layout':
            bi_mode = 'inline'
        elif bi_mode == 'line_by_line':
            bi_mode = 'newline'
        if bi_mode not in ('none', 'inline', 'newline'):
            bi_mode = 'none'
        # Keep document line/paragraph structure stable.
        if bi_mode == 'newline':
            bi_mode = 'none'
        # Strict layout policy: do not create new lines/paragraphs in DOCX.
        if bi_mode == 'newline':
            bi_mode = 'none'
        bi_delim = self._normalize_bilingual_delimiter(bilingual_delimiter) if bi_mode == 'inline' else '|'

        # Output path in downloads (consistent with other document outputs)
        base = os.path.splitext(os.path.basename(file_path))[0]
        out_name = f"{base}_translated_{uuid.uuid4().hex[:8]}.pdf"
        out_path = os.path.join(self.download_folder, out_name)

        # Simple cache to avoid repeated API calls for identical lines.
        cache = {}

        def _should_translate(s: str) -> bool:
            if s is None:
                return False
            if not str(s).strip():
                return False
            core = str(s).strip()
            if re.fullmatch(r"[\d\W_]+", core, flags=re.UNICODE):
                return False
            if skip_vietnamese_source and _is_probably_vietnamese(core):
                return False
            return True

        def _translate_preserve_ws(s: str) -> str:
            src = "" if s is None else str(s)
            m = re.match(r"^(\s*)(.*?)(\s*)$", src, flags=re.DOTALL)
            lead, core, tail = (m.group(1), m.group(2), m.group(3)) if m else ("", src, "")
            if not _should_translate(core):
                return src
            if core in cache:
                return f"{lead}{cache[core]}{tail}"
            dst = self._translate_with_retry(core, target_lang, context='document_pdf')
            dst = "" if dst is None else str(dst)
            cache[core] = dst
            return f"{lead}{dst}{tail}"

        def _int_color_to_rgb01(color_int: int):
            try:
                c = int(color_int)
            except Exception:
                c = 0
            r = ((c >> 16) & 255) / 255.0
            g = ((c >> 8) & 255) / 255.0
            b = (c & 255) / 255.0
            return (r, g, b)

        # ── Unicode font resolution (supports Vietnamese/CJK diacritics) ──
        # Map font families to system font files with bold/italic variants.
        _FONT_DIR = os.environ.get("FONT_DIR", "")
        if not _FONT_DIR or not os.path.isdir(_FONT_DIR):
            # Windows
            if os.path.isdir(r"C:\Windows\Fonts"):
                _FONT_DIR = r"C:\Windows\Fonts"
            # Linux / Docker
            elif os.path.isdir("/usr/share/fonts/truetype"):
                _FONT_DIR = "/usr/share/fonts/truetype"
            else:
                _FONT_DIR = ""

        # (regular, bold, italic, bold-italic) file names
        _FONT_FAMILIES = {
            "sans": ("arial.ttf", "arialbd.ttf", "ariali.ttf", "arialbi.ttf"),
            "serif": ("times.ttf", "timesbd.ttf", "timesi.ttf", "timesbi.ttf"),
            "mono": ("cour.ttf", "courbd.ttf", "couri.ttf", "courbi.ttf"),
        }
        # Linux fallback names
        _FONT_FAMILIES_LINUX = {
            "sans": ("DejaVuSans.ttf", "DejaVuSans-Bold.ttf", "DejaVuSans-Oblique.ttf", "DejaVuSans-BoldOblique.ttf"),
            "serif": ("DejaVuSerif.ttf", "DejaVuSerif-Bold.ttf", "DejaVuSerif-Italic.ttf", "DejaVuSerif-BoldItalic.ttf"),
            "mono": ("DejaVuSansMono.ttf", "DejaVuSansMono-Bold.ttf", "DejaVuSansMono-Oblique.ttf", "DejaVuSansMono-BoldOblique.ttf"),
        }

        def _find_font_file(family_key, variant_idx):
            """Find a system font file. variant_idx: 0=regular, 1=bold, 2=italic, 3=bolditalic."""
            for families in (_FONT_FAMILIES, _FONT_FAMILIES_LINUX):
                names = families.get(family_key, families["sans"])
                fname = names[variant_idx]
                # Direct path
                path = os.path.join(_FONT_DIR, fname)
                if os.path.isfile(path):
                    return path
                # Search subdirectories (Linux: /usr/share/fonts/truetype/dejavu/)
                if _FONT_DIR:
                    for root, _dirs, files in os.walk(_FONT_DIR):
                        if fname in files:
                            return os.path.join(root, fname)
            return None

        # Cache resolved font paths
        _font_cache = {}

        def _resolve_font(style_span: dict):
            """Resolve a Unicode-capable font file + internal name for a span's style.

            Returns (fontname_str, fontfile_path_or_None, family_key).
            If fontfile is None, falls back to Base-14 names.
            """
            name = str((style_span or {}).get("font") or "").lower()
            flags = 0
            try:
                flags = int((style_span or {}).get("flags") or 0)
            except Exception:
                flags = 0

            is_bold = ("bold" in name) or bool(flags & 16)
            is_italic = ("italic" in name) or ("oblique" in name) or bool(flags & 2)

            # Classify family
            family = "sans"
            if "times" in name or "tiro" in name or "serif" in name or "georgia" in name:
                family = "serif"
            elif "cour" in name or "mono" in name or "consol" in name:
                family = "mono"

            # variant index: 0=regular, 1=bold, 2=italic, 3=bolditalic
            variant = 0
            if is_bold and is_italic:
                variant = 3
            elif is_bold:
                variant = 1
            elif is_italic:
                variant = 2

            cache_key = (family, variant)
            if cache_key in _font_cache:
                return _font_cache[cache_key]

            fontfile = _find_font_file(family, variant)
            if fontfile:
                # Use a unique internal name per variant to avoid collisions
                internal_name = f"F{family[0]}{variant}"
                result = (internal_name, fontfile, family)
            else:
                # Fallback to Base-14 (no Vietnamese support, but won't crash)
                base14_map = {
                    ("sans", 0): "Helvetica", ("sans", 1): "Helvetica-Bold",
                    ("sans", 2): "Helvetica-Oblique", ("sans", 3): "Helvetica-BoldOblique",
                    ("serif", 0): "Times-Roman", ("serif", 1): "Times-Bold",
                    ("serif", 2): "Times-Italic", ("serif", 3): "Times-BoldItalic",
                    ("mono", 0): "Courier", ("mono", 1): "Courier-Bold",
                    ("mono", 2): "Courier-Oblique", ("mono", 3): "Courier-BoldOblique",
                }
                result = (base14_map.get(cache_key, "Helvetica"), None, family)

            _font_cache[cache_key] = result
            return result

        def _insert_text_fit(page, rect, text, *, fontname, fontfile, fontsize, color):
            """Insert text at rect position. Try textbox first; fall back to insert_text."""
            fs0 = int(max(4, round(float(fontsize))))
            # Try insert_textbox (wraps text within bounds)
            for fs in range(fs0, 3, -1):
                try:
                    kwargs = dict(fontsize=fs, fontname=fontname, color=color, align=0)
                    if fontfile:
                        kwargs["fontfile"] = fontfile
                    rc = page.insert_textbox(rect, text, **kwargs)
                    if rc >= 0:
                        return True
                except Exception:
                    try:
                        rc = page.insert_textbox(
                            rect, text, fontsize=fs, fontname="Helvetica",
                            color=color, align=0,
                        )
                        if rc >= 0:
                            return True
                    except Exception:
                        continue
            # Fallback: point-based insert_text (won't clip but may overflow to the right)
            try:
                baseline_y = rect.y1 - 1
                kwargs = dict(fontsize=fs0, fontname=fontname, color=color)
                if fontfile:
                    kwargs["fontfile"] = fontfile
                page.insert_text(fitz.Point(rect.x0, baseline_y), text, **kwargs)
                return True
            except Exception:
                try:
                    page.insert_text(
                        fitz.Point(rect.x0, rect.y1 - 1), text,
                        fontsize=fs0, fontname="Helvetica", color=color,
                    )
                    return True
                except Exception:
                    return False

        doc = fitz.open(file_path)
        try:
            total_pages = int(getattr(doc, "page_count", 0) or len(doc))
            if total_pages <= 0:
                raise RuntimeError("Empty PDF")

            # ── For bilingual newline: build new PDF with image backgrounds ──
            _bi_out_doc = None
            if bi_mode == 'newline':
                _bi_out_doc = fitz.open()

            for page_index in range(total_pages):
                page = doc.load_page(page_index)
                src_page = page
                text_dict = page.get_text("dict")
                page_w = page.rect.width
                page_h = page.rect.height

                # Collect line items first to allow a deterministic processing order.
                items = []  # (line_rect, span_rects, src_text, style_span)
                for block in (text_dict.get("blocks") or []):
                    if block.get("type") != 0:
                        continue
                    for line in (block.get("lines") or []):
                        spans = line.get("spans") or []
                        if not spans:
                            continue
                        src = "".join((sp.get("text") or "") for sp in spans)
                        if not _should_translate(src):
                            continue
                        # Style: take the first non-empty span in this line.
                        style = None
                        for sp in spans:
                            if (sp.get("text") or "").strip():
                                style = sp
                                break
                        if not style:
                            continue
                        try:
                            line_rect = fitz.Rect(line.get("bbox"))
                        except Exception:
                            continue

                        # Ignore tiny boxes (often artifacts)
                        if line_rect.width < 2 or line_rect.height < 2:
                            continue

                        span_rects = []
                        for sp in spans:
                            if not (sp.get("text") or "").strip():
                                continue
                            try:
                                r = fitz.Rect(sp.get("bbox"))
                            except Exception:
                                continue
                            if r.width < 1 or r.height < 1:
                                continue
                            span_rects.append(r)

                        if not span_rects:
                            continue

                        items.append((line_rect, span_rects, src, style))

                if progress_callback:
                    base_pct = int(5 + (page_index / max(1, total_pages)) * 90)
                    progress_callback(base_pct, f"PDF: scanning page {page_index+1}/{total_pages}")

                if not items:
                    continue

                # ── Bilingual newline: render page as image, overlay original + translation ──
                if bi_mode == 'newline':
                    # Render original page as image background (preserves borders/logos/layout)
                    pix = src_page.get_pixmap(dpi=200)
                    new_page = _bi_out_doc.new_page(width=page_w, height=page_h)
                    new_page.insert_image(fitz.Rect(0, 0, page_w, page_h), pixmap=pix)

                    RIGHT_MARGIN = 30
                    for idx, (rect, span_rects, src, style) in enumerate(items, start=1):
                        # Translate
                        try:
                            dst = _translate_preserve_ws(src)
                        except ProviderRateLimitError:
                            raise
                        except Exception:
                            dst = ""

                        fontname, fontfile, _family = _resolve_font(style)
                        fontsize = style.get("size") or 10
                        color = _int_color_to_rgb01(style.get("color") or 0)

                        # White-out span rects on the image
                        for sr in span_rects:
                            new_page.draw_rect(sr, color=None, fill=(1, 1, 1), overlay=True, width=0)

                        # Reinsert original text at its exact position
                        _insert_text_fit(new_page, rect, src.strip(),
                                         fontname=fontname, fontfile=fontfile,
                                         fontsize=fontsize, color=color)

                        # Insert translation below in blue italic
                        if str(dst).strip():
                            fn_i, ff_i, _ = _resolve_font({"font": style.get("font", "arial"), "flags": 2})
                            trans_fs = max(5, fontsize - 2)
                            trans_color = (0.0, 0.10, 0.65)
                            trans_h = max(trans_fs + 2, rect.height * 0.85)
                            trans_rect = fitz.Rect(
                                rect.x0,
                                rect.y1 + 1,
                                max(rect.x1, page_w - RIGHT_MARGIN),
                                rect.y1 + 1 + trans_h
                            )
                            new_page.draw_rect(trans_rect, color=None, fill=(1, 1, 1), overlay=True, width=0)
                            _insert_text_fit(new_page, trans_rect, str(dst).strip(),
                                             fontname=fn_i, fontfile=ff_i,
                                             fontsize=trans_fs, color=trans_color)

                        if progress_callback and idx % 10 == 0:
                            pct = int(5 + ((page_index + (idx / max(1, len(items)))) / max(1, total_pages)) * 90)
                            progress_callback(min(98, pct), f"PDF song ngữ: dịch trang {page_index+1}/{total_pages}")

                    continue  # skip normal redact+insert for this page


                # ── Remove original text using redaction (preserves table borders) ──
                for _line_rect, span_rects, _src, _style in items:
                    for r in span_rects:
                        try:
                            page.add_redact_annot(r, fill=False)
                        except Exception:
                            pass

                try:
                    page.apply_redactions(
                        images=fitz.PDF_REDACT_IMAGE_NONE,
                        graphics=fitz.PDF_REDACT_LINE_ART_NONE,
                    )
                except Exception:
                    for _line_rect, span_rects, _src, _style in items:
                        for r in span_rects:
                            try:
                                page.draw_rect(r, color=None, fill=(1, 1, 1), overlay=True, width=0)
                            except Exception:
                                pass

                # ── Insert text back ──
                for idx, (rect, _span_rects, src, style) in enumerate(items, start=1):
                    try:
                        dst = _translate_preserve_ws(src)
                    except ProviderRateLimitError:
                        raise
                    except Exception:
                        dst = ""

                    if not str(dst).strip():
                        continue

                    fontname, fontfile, _family = _resolve_font(style)
                    fontsize = style.get("size") or 10
                    color = _int_color_to_rgb01(style.get("color") or 0)

                    if bi_mode == 'inline':
                        # Song ngữ liền kề: "Original <delim> Translated"
                        # Use insert_text (point-based) so it doesn't clip in narrow cells.
                        display_text = self._join_inline_bilingual(src.strip(), str(dst).strip(), bi_delim)
                        # Baseline position: near the bottom of the line rect
                        baseline_y = rect.y1 - 1
                        best_fs = int(max(4, round(float(fontsize))))
                        kwargs = dict(fontname=fontname, color=color)
                        if fontfile:
                            kwargs["fontfile"] = fontfile
                        try:
                            page.insert_text(fitz.Point(rect.x0, baseline_y), display_text, fontsize=best_fs, **kwargs)
                        except Exception:
                            try:
                                page.insert_text(
                                    fitz.Point(rect.x0, baseline_y), display_text,
                                    fontsize=best_fs, fontname="Helvetica", color=color,
                                )
                            except Exception:
                                pass
                    else:
                        # Normal mode: only translated text, use textbox for fit
                        _insert_text_fit(page, rect, str(dst),
                                         fontname=fontname, fontfile=fontfile,
                                         fontsize=fontsize, color=color)

                    if progress_callback and idx % 40 == 0:
                        pct = int(5 + ((page_index + (idx / max(1, len(items)))) / max(1, total_pages)) * 90)
                        progress_callback(min(98, pct), f"PDF: translating page {page_index+1}/{total_pages} ({idx}/{len(items)} lines)")

            if _bi_out_doc is not None:
                _bi_out_doc.save(out_path, garbage=4, deflate=True)
                _bi_out_doc.close()
            else:
                doc.save(out_path, garbage=4, deflate=True)
        finally:
            try:
                doc.close()
            except Exception:
                pass

        if progress_callback:
            progress_callback(100, "PDF: completed")
        return out_path

    @staticmethod
    def _normalize_bilingual_delimiter(delimiter):
        """Normalize a user-provided delimiter for inline bilingual output.

        Returns a short string (no surrounding spaces). Defaults to '|'.
        """
        d = (delimiter or '').strip()
        if not d:
            return '|'
        # Prevent pathological inputs from breaking layout.
        if len(d) > 10:
            d = d[:10]
        return d

    def _join_inline_bilingual(self, src, dst, delimiter):
        d = self._normalize_bilingual_delimiter(delimiter)
        s = (src or '').strip()
        t = (dst or '').strip()
        # If translation is empty, do not append a dangling delimiter.
        if not t:
            return src or ''
        if not s:
            return dst or ''
        return f"{src} {d} {dst}"

    def _process_docx(self, file_path, target_lang, progress_callback=None, *, ocr_images=False, ocr_langs=None, ocr_mode=None, bilingual_mode=None, bilingual_delimiter=None):
        """Translate DOCX while preserving original formatting, layout, images.
        
        Bilingual modes:
          - none: normal translation (replace original with translation)
          - inline: song ngữ liền kề (Original | Translated in same paragraph)
          - newline: song ngữ xuống dòng (keep original, add translated paragraph below)
          - preserve_layout: alias for 'inline' mode (dịch song ngữ liền kề, giữ layout)
          - line_by_line: alias for 'newline' mode (dịch song ngữ xuống dòng)
        """
        # Modify original document in-place so styles/images/relationships are preserved
        doc = docx.Document(file_path)

        # ── Ensure table borders are explicitly set in XML ──
        # python-docx preserves existing XML, but tables whose borders rely
        # solely on a style definition can lose their borders when runs/paragraphs
        # are modified. Stamp explicit <w:tblBorders> and <w:tcBorders> from
        # the style into the table/cell properties so they survive the save.
        try:
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn as _qn
            import copy as _copy

            # Table borders are preserved as-is from the original document.
            # We do NOT force borders on tables, as this would alter borderless tables.
            pass
        except Exception:
            pass

        api_only = str(os.getenv('AI_DISABLE_FALLBACK', '0')).strip().lower() in ('1', 'true', 'yes', 'on')

        # Bilingual mode:
        # - None / 'none': normal (replace original with translation)
        # - 'inline':  song ngữ liền kề — "Original | Translated" in same paragraph
        # - 'newline': song ngữ xuống dòng — keep original, add translated paragraph below
        # - 'preserve_layout': dịch song ngữ liền kề - giữ layout gốc (tương đương 'inline')
        # - 'line_by_line': dịch song ngữ xuống dòng - text gốc trên, dịch dưới (tương đương 'newline')
        bi_mode = (str(bilingual_mode).strip().lower() if bilingual_mode else 'none')
        if bi_mode not in ('none', 'inline', 'newline', 'preserve_layout', 'line_by_line'):
            bi_mode = 'none'
        
        # Normalize mode names: preserve_layout -> inline, line_by_line -> newline
        if bi_mode == 'preserve_layout':
            bi_mode = 'inline'
        elif bi_mode == 'line_by_line':
            bi_mode = 'newline'

        # OCR mode for embedded images in DOCX:
        # - image: replace embedded image bytes with overlay-rendered translation (keep design)
        # - text:  replace current image with translated text at that image location
        # - both:  keep translated image + also add translated text paragraph
        # - auto:  smart pick per image (prefers 'text' when Tesseract not available)
        mode = (str(ocr_mode).strip().lower() if ocr_mode else 'image')
        if mode not in ('image', 'text', 'both', 'auto'):
            mode = 'auto'

        def _auto_pick_mode(ocr_text, translated_text, ai_recommended_mode=None):
            """Pick OCR output mode per embedded image (AUTO mode).

            Returns 'text' or 'image':
              'text'  -> giữ ảnh gốc + chèn bản dịch bên dưới
              'image' -> chồng bản dịch lên ảnh (overlay)

            Logic đơn giản:
              - Ảnh có nhiều chữ (bài đọc, văn bản) -> 'text'
              - Ảnh kiểu banner / poster / quảng cáo  -> 'image'
            """
            try:
                raw = (ocr_text or '').strip()
                if not raw:
                    return 'text'

                char_count = len(raw)
                words = re.findall(r'\w+', raw, flags=re.UNICODE)
                word_count = len(words)

                # ── Prose: nhiều chữ -> luôn 'text' (giữ ảnh gốc) ──
                if char_count >= 120 or word_count >= 25:
                    print(f"  [MODE] Prose detected (chars={char_count}, words={word_count}), AI={ai_recommended_mode} -> text")
                    return 'text'

                # ── Banner detection ──
                lines = [ln.strip() for ln in raw.splitlines() if ln.strip()]
                low_raw = raw.lower()
                low_trans = (translated_text or '').lower()
                promo_keywords = (
                    'sale', 'discount', 'offer', 'book now', 'vacation',
                    'summer', 'up to', '% off', 'promo', 'hotline',
                    'free', 'limited', 'special', 'deal', 'subscribe',
                )
                has_promo = any(k in low_raw or k in low_trans for k in promo_keywords)

                alpha_chars = [ch for ch in raw if ch.isalpha()]
                upper_ratio = (
                    sum(1 for ch in alpha_chars if ch == ch.upper()) / max(1, len(alpha_chars))
                    if alpha_chars else 0.0
                )
                line_word_counts = [len(re.findall(r'\w+', ln, flags=re.UNICODE)) for ln in lines] if lines else [0]
                avg_wpl = (sum(line_word_counts) / len(line_word_counts)) if line_word_counts else 0.0
                short_lines = sum(1 for c in line_word_counts if c <= 3)

                looks_banner = (
                    has_promo or
                    (upper_ratio >= 0.50 and avg_wpl <= 4) or
                    (short_lines >= 3 and avg_wpl <= 3)
                )

                ai_mode = (ai_recommended_mode or '').lower()

                if looks_banner or ai_mode == 'image':
                    final = 'image'
                else:
                    final = 'text'

                print(
                    f"  [MODE] AI={ai_mode}, banner={looks_banner}, "
                    f"chars={char_count}, words={word_count}, upper={upper_ratio:.2f} -> {final}"
                )
                return final
            except Exception:
                return 'text'

        def iter_all_paragraphs(document):
            paras = []
            try:
                paras.extend(list(document.paragraphs))
            except Exception:
                pass
            try:
                for table in document.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            paras.extend(list(cell.paragraphs))
            except Exception:
                pass
            try:
                for section in document.sections:
                    paras.extend(list(section.header.paragraphs))
                    paras.extend(list(section.footer.paragraphs))
            except Exception:
                pass
            return paras

        def paragraph_image_rids(paragraph):
            # Return relationship ids (rIdX) for images embedded in this paragraph.
            rids = []
            try:
                runs = list(paragraph.runs)
            except Exception:
                runs = []
            if not runs:
                return rids

            # Use a namespace-agnostic xpath for blips (image references)
            rel_attr = '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed'
            for run in runs:
                try:
                    blips = run._element.xpath('.//*[local-name()="blip"]')
                except Exception:
                    blips = []
                for blip in blips:
                    try:
                        rid = blip.get(rel_attr)
                    except Exception:
                        rid = None
                    if rid:
                        rids.append(rid)
            # Preserve order but de-dup
            seen = set()
            out = []
            for rid in rids:
                if rid in seen:
                    continue
                seen.add(rid)
                out.append(rid)
            return out

        def rid_to_image_part(paragraph, rid):
            try:
                part = paragraph.part
                # python-docx keeps a mapping of related parts by rId
                related = getattr(part, 'related_parts', None)
                if isinstance(related, dict) and rid in related:
                    return related[rid]
            except Exception:
                pass
            try:
                # Fallback: via relationship object
                rels = getattr(paragraph.part, 'rels', None)
                if rels and rid in rels:
                    return rels[rid].target_part
            except Exception:
                pass
            return None

        def _collect_header_footer_image_partnames(document):
            protected = set()
            try:
                for section in document.sections:
                    for hf in (section.header, section.footer):
                        part = getattr(hf, 'part', None)
                        related = getattr(part, 'related_parts', None)
                        if not isinstance(related, dict):
                            continue
                        for _rid, target in related.items():
                            try:
                                ct = str(getattr(target, 'content_type', '') or '').lower()
                                if not ct.startswith('image/'):
                                    continue
                                pn = str(getattr(target, 'partname', '') or '').lstrip('/')
                                if pn:
                                    protected.add(pn)
                            except Exception:
                                continue
            except Exception:
                pass
            return protected

        def replace_image_with_text(paragraph, rid, translated_text):
            """Replace a specific embedded image (by rid) in a paragraph with translated text."""
            txt = (translated_text or '').strip()
            if not txt:
                return False

            rel_attr = '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed'
            replaced = False
            try:
                runs = list(paragraph.runs)
            except Exception:
                runs = []

            for run in runs:
                try:
                    blips = run._element.xpath('.//*[local-name()="blip"]')
                except Exception:
                    blips = []
                if not blips:
                    continue

                has_target = False
                for blip in blips:
                    try:
                        if blip.get(rel_attr) == rid:
                            has_target = True
                            break
                    except Exception:
                        continue
                if not has_target:
                    continue

                try:
                    drawings = run._element.xpath('./*[local-name()="drawing"]')
                    for dr in drawings:
                        parent = dr.getparent()
                        if parent is not None:
                            parent.remove(dr)
                except Exception:
                    pass
                # Remove residual text from the image run; we add a clean run below.
                run.text = ""
                replaced = True
                break

            # Insert a clean run with normalized paragraph settings to avoid stretched spacing.
            try:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            except Exception:
                pass

            try:
                new_run = paragraph.add_run(txt)
                _ = new_run
                replaced = True
            except Exception:
                pass

            if not replaced:
                try:
                    paragraph.add_run(txt)
                    replaced = True
                except Exception:
                    replaced = False
            return replaced

        def _normalize_ocr_text_for_docx(text):
            """Normalize OCR translated text for readable DOCX insertion.

            - remove noisy tiny lines
            - collapse excessive whitespace
            - reflow short broken lines into normal prose
            """
            raw = (text or '').replace('\r\n', '\n').replace('\r', '\n')
            if not raw.strip():
                return ''

            cleaned_lines = []
            for ln in raw.split('\n'):
                ln2 = re.sub(r'\s+', ' ', (ln or '').strip())
                if not ln2:
                    continue
                if len(ln2) <= 1 and not re.search(r'[0-9]', ln2):
                    continue
                cleaned_lines.append(ln2)

            if not cleaned_lines:
                return ''

            out_parts = []
            cur = ''
            for ln in cleaned_lines:
                if not cur:
                    cur = ln
                    continue

                end_punct = cur.endswith(('.', '!', '?', ':', ';'))
                starts_bullet = bool(re.match(r'^(\-|\*|\d+[\.)])\s+', ln))
                if end_punct or starts_bullet:
                    out_parts.append(cur)
                    cur = ln
                else:
                    cur = f"{cur} {ln}".strip()

            if cur:
                out_parts.append(cur)

            normalized = '\n'.join(out_parts)
            return normalized.strip()

        DB_IDENTIFIER_MAP = {
            "ma_khach_hang": "customer_id",
            "ngay_ban": "sale_date",
        }

        def _apply_db_identifier_map(text: str) -> str:
            out = text or ""
            for src, dst in DB_IDENTIFIER_MAP.items():
                out = re.sub(rf"\b{re.escape(src)}\b", dst, out, flags=re.IGNORECASE)
            return out

        def _cleanup_translated_text(text: str) -> str:
            """Apply minimal lexical fixes without changing whitespace/line layout."""
            out = "" if text is None else str(text)
            out = re.sub(r"\bNAMEBUILDING\b", "NAME BUILDING", out, flags=re.IGNORECASE)
            out = _apply_db_identifier_map(out)
            out = re.sub(r"\bNot\s+nul\b", "Not null", out, flags=re.IGNORECASE)
            out = re.sub(r"\bInfo\s+tin\s+basic\b", "Basic Information", out, flags=re.IGNORECASE)
            out = re.sub(r"\bWhere\s+the\s+topic\s+is\s+applied\b", "Application of the project", out, flags=re.IGNORECASE)
            out = re.sub(r"\bDevelopment\s+direction\s*:\s*Is\s+there\b", "Development direction: Yes", out, flags=re.IGNORECASE)
            out = re.sub(r"\bSTUDENT\s+ID\b", "Student ID", out, flags=re.IGNORECASE)
            out = re.sub(r"\bData\s+types\b", "Data Types", out, flags=re.IGNORECASE)
            out = re.sub(r"\bKHOA\s*CÔNG\s*NGHỆ\s*THÔNG\s*TIN\b", "FACULTY OF INFORMATION TECHNOLOGY", out, flags=re.IGNORECASE)
            out = re.sub(r"Gửi\s+lại\s+phiếu\s+đăng\s+ký\s+qua\s+Email\s*:", "Resubmit the registration form via Email:", out, flags=re.IGNORECASE)
            return out

        def _translate_preserve_exact_lines(text):
            """Translate line-by-line to avoid line merge/split in DOCX."""
            raw = text or ""
            if not raw:
                return raw
            parts = re.split(r"(\r\n|\r|\n)", raw)
            out = []
            for part in parts:
                if part in ("\r\n", "\r", "\n"):
                    out.append(part)
                    continue
                if not part.strip():
                    out.append(part)
                    continue
                m = re.match(r"^(\s*)(.*?)(\s*)$", part, flags=re.DOTALL)
                if m:
                    lead, core, tail = m.group(1), m.group(2), m.group(3)
                else:
                    lead, core, tail = "", part, ""
                if not core.strip():
                    out.append(part)
                    continue
                translated_core = _translate_preserve_form_leaders(core)
                out.append(f"{lead}{translated_core}{tail}")
            return "".join(out)

        def image_part_ext(image_part):
            # Try best-effort extension resolution
            try:
                partname = str(getattr(image_part, 'partname', '') or '')
                base = os.path.basename(partname)
                ext = os.path.splitext(base)[1].lower()
                if ext:
                    return ext
            except Exception:
                pass
            try:
                ct = str(getattr(image_part, 'content_type', '') or '').lower()
                mapping = {
                    'image/png': '.png',
                    'image/jpeg': '.jpg',
                    'image/jpg': '.jpg',
                    'image/gif': '.gif',
                    'image/bmp': '.bmp',
                    'image/tiff': '.tif',
                    'image/webp': '.webp',
                }
                return mapping.get(ct, '.png')
            except Exception:
                return '.png'

        def _overlay_bytes_to_original_format(png_bytes: bytes, desired_ext: str) -> bytes:
            """Convert PNG bytes (rendered overlay) to match the original image extension."""
            desired_ext = (desired_ext or '.png').lower()
            try:
                from PIL import Image
            except Exception:
                # If Pillow is not available, return PNG bytes as-is.
                return png_bytes

            fmt_map = {
                '.png': 'PNG',
                '.jpg': 'JPEG',
                '.jpeg': 'JPEG',
                '.bmp': 'BMP',
                '.tif': 'TIFF',
                '.tiff': 'TIFF',
                '.webp': 'WEBP',
                '.gif': 'PNG',  # avoid animated GIF issues
            }
            out_fmt = fmt_map.get(desired_ext, 'PNG')
            try:
                img = Image.open(io.BytesIO(png_bytes))
                if out_fmt in ('JPEG', 'BMP', 'TIFF'):
                    if img.mode not in ('RGB', 'L'):
                        img = img.convert('RGB')
                buf = io.BytesIO()
                img.save(buf, format=out_fmt)
                return buf.getvalue()
            except Exception:
                return png_bytes

        def _apply_translation_to_runs(paragraph, translated_text):
            """Apply translated text preserving structural runs (leaders, tabs).

            Runs that contain ONLY dots/underscores/dashes/tabs (no word chars)
            are kept untouched.  Only "content" runs are replaced.
            """
            runs = list(paragraph.runs)
            if not runs:
                paragraph.add_run(translated_text or "")
                return

            # Classify each run as content vs structural
            content_indices = []
            structural_indices = []
            for i, run in enumerate(runs):
                rt = run.text or ""
                if _is_structural_text(rt):
                    structural_indices.append(i)
                else:
                    content_indices.append(i)

            if not content_indices:
                # No translatable content — keep everything as-is
                return

            if structural_indices:
                # Has structural runs (leaders/tabs): only replace content runs,
                # leave structural runs untouched.
                # Translated text should NOT include the leaders (they stay in place),
                # so we only pass the content portion to translation upstream.
                primary = content_indices[0]
                runs[primary].text = translated_text or ""
                for i in content_indices[1:]:
                    runs[i].text = ""
                # structural_indices are NOT touched → leaders/tabs preserved
            else:
                # No structural runs: put everything in primary run
                primary = content_indices[0]
                runs[primary].text = translated_text or ""
                for i, r in enumerate(runs):
                    if i != primary:
                        r.text = ""

        def _get_run_format_key(run):
            """Return a hashable key representing this run's formatting (rPr XML)."""
            from docx.oxml.ns import qn as _qn
            try:
                from lxml import etree
                rPr = run._element.find(_qn('w:rPr'))
                if rPr is not None:
                    # Use plain tostring (c14n2 fails on OOXML namespaces)
                    return etree.tostring(rPr)
                return b''
            except Exception:
                return b''

        def _group_runs_by_format(runs):
            """Group consecutive runs with the same formatting.

            Returns list of (format_key, [run_indices]).
            Whitespace-only runs are merged into the preceding group.
            Structural runs (only dots/underscores/tabs) get their own group
            so they are never merged with translatable content.
            """
            groups = []
            for i, run in enumerate(runs):
                text = run.text or ""
                fmt = _get_run_format_key(run)

                # Structural runs (leaders, tabs) → always separate group
                # so they are never accidentally sent to the translator.
                if text.strip() and _is_structural_text(text):
                    groups.append((b'__structural__' + fmt, [i]))
                    continue

                if not text.strip():
                    # Whitespace-only: attach to current group if exists
                    if groups:
                        groups[-1][1].append(i)
                    else:
                        groups.append((fmt, [i]))
                    continue

                if groups and groups[-1][0] == fmt:
                    groups[-1][1].append(i)
                else:
                    groups.append((fmt, [i]))
            return groups

        def _translate_format_groups(paragraph, translate_fn):
            """Translate a paragraph by grouping runs with same formatting.

            Each format-group is translated independently so run-level formatting
            (bold, italic, color, font, size) is perfectly preserved.
            Structural runs (dots, underscores, tabs) are never translated.
            """
            from docx.oxml.ns import qn as _qn
            runs = list(paragraph.runs)
            if not runs:
                return

            original_texts = [(r.text or "") for r in runs]
            paragraph_text = "".join(original_texts)
            if not paragraph_text.strip():
                return

            # Check if any run is structural (leaders/tabs).
            has_structural_runs = any(
                (original_texts[i] or "").strip() and _is_structural_text(original_texts[i])
                for i in range(len(runs))
            )

            groups = _group_runs_by_format(runs)

            # If only one group AND no structural runs → translate entire paragraph
            if len(groups) <= 1 and not has_structural_runs:
                translated = translate_fn(paragraph_text)
                _apply_translation_to_runs(paragraph, translated)
                return

            # If only one group BUT has structural runs → use leader-aware translate
            # on the whole text, then write back preserving structural runs.
            if len(groups) <= 1 and has_structural_runs:
                translated = translate_fn(paragraph_text)
                _apply_translation_to_runs(paragraph, translated)
                return

            # Multiple format groups → translate each group separately
            for fmt_key, indices in groups:
                group_text = "".join(original_texts[i] for i in indices)
                if not group_text.strip():
                    continue

                # Skip structural groups entirely (leaders, tabs, pure punctuation)
                if _is_structural_text(group_text):
                    continue

                try:
                    translated_group = translate_fn(group_text)
                except ProviderRateLimitError:
                    raise
                except Exception as e:
                    print(f"Format-group translation failed: {e}")
                    if api_only:
                        raise
                    translated_group = group_text

                # Write into first non-empty run of the group; clear others
                written = False
                for i in indices:
                    run = runs[i]
                    if not written and (original_texts[i] or "").strip():
                        run.text = translated_group or ""
                        written = True
                    else:
                        run.text = ""

                if not written and indices:
                    runs[indices[0]].text = translated_group or ""
                    for i in indices[1:]:
                        runs[i].text = ""

        # ── Helper: insert a new paragraph right after `ref_para` in the document body ──
        def _insert_paragraph_after(ref_para, text, italic=True):
            """Insert a translated paragraph immediately after ref_para.

            Copies ALL paragraph properties from the source (except numbering
            to avoid duplicate list bullets), preserving the original format.
            Also strips distributed alignment and expanded character spacing
            so the translated paragraph renders normally.
            """
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn as _qn
            import copy as _copy

            new_p = OxmlElement('w:p')

            # Deep-copy the entire paragraph properties block from the source.
            try:
                pPr_src = ref_para._element.find(_qn('w:pPr'))
                if pPr_src is not None:
                    new_pPr = _copy.deepcopy(pPr_src)

                    # ── Handle numbering ──
                    # Remove numPr to prevent duplicate bullets/numbering.
                    # Before removing, capture the source paragraph's effective
                    # indentation so we can set it explicitly on the new paragraph
                    # (numbered paragraphs often derive their indent from the
                    # numbering definition, not from w:ind).
                    numPr = new_pPr.find(_qn('w:numPr'))
                    if numPr is not None:
                        # Read the source paragraph's effective indent BEFORE removing numPr.
                        src_left = None
                        src_hanging = None
                        src_firstLine = None
                        try:
                            fmt = ref_para.paragraph_format
                            src_left = fmt.left_indent          # EMU int or None
                            src_firstLine = fmt.first_line_indent  # EMU int or None (negative = hanging)
                        except Exception:
                            pass
                        # Also try to read from the original w:ind XML as fallback.
                        if src_left is None:
                            try:
                                orig_ind = pPr_src.find(_qn('w:ind'))
                                if orig_ind is not None:
                                    l_val = orig_ind.get(_qn('w:left')) or orig_ind.get(_qn('w:start'))
                                    if l_val:
                                        src_left = int(l_val)
                                    h_val = orig_ind.get(_qn('w:hanging'))
                                    if h_val:
                                        src_hanging = int(h_val)
                                    fl_val = orig_ind.get(_qn('w:firstLine'))
                                    if fl_val:
                                        src_firstLine = int(fl_val)
                            except Exception:
                                pass

                        new_pPr.remove(numPr)

                        # Ensure the new paragraph has an explicit w:ind so its
                        # indentation matches the source paragraph visually.
                        ind = new_pPr.find(_qn('w:ind'))
                        if ind is None and src_left is not None:
                            ind = OxmlElement('w:ind')
                            ind.set(_qn('w:left'), str(src_left))
                            if src_hanging is not None:
                                ind.set(_qn('w:hanging'), str(src_hanging))
                            elif src_firstLine is not None:
                                if src_firstLine < 0:
                                    ind.set(_qn('w:hanging'), str(abs(src_firstLine)))
                                elif src_firstLine > 0:
                                    ind.set(_qn('w:firstLine'), str(src_firstLine))
                            new_pPr.append(ind)

                    # ── Strip distributed / thai-distribute alignment ──
                    # Distributed alignment spreads characters across the full
                    # line width — this looks wrong for translated text which
                    # typically has a different character count.
                    jc = new_pPr.find(_qn('w:jc'))
                    if jc is not None:
                        jc_val = jc.get(_qn('w:val'), '')
                        if jc_val in ('distribute', 'thai-distribute'):
                            new_pPr.remove(jc)

                    new_p.insert(0, new_pPr)
            except Exception:
                pass

            run_el = OxmlElement('w:r')
            rPr = OxmlElement('w:rPr')

            # Inherit run formatting from the first non-empty run of source.
            # Prefer a run with explicit color, fallback to first non-empty.
            # Strip expanded character spacing (w:spacing) so translated text
            # is not rendered with artificially wide letter spacing.
            try:
                src_runs = ref_para._element.findall('.//' + _qn('w:r'))
                chosen_rPr = None
                # 1) Prefer a run that has explicit color formatting
                for r in src_runs:
                    try:
                        t = r.find(_qn('w:t'))
                        if t is None or not (t.text or '').strip():
                            continue
                        rpr = r.find(_qn('w:rPr'))
                        if rpr is None:
                            continue
                        if rpr.find(_qn('w:color')) is not None:
                            chosen_rPr = rpr
                            break
                    except Exception:
                        continue
                # 2) Fallback: first non-empty run
                if chosen_rPr is None:
                    for r in src_runs:
                        try:
                            t = r.find(_qn('w:t'))
                            if t is not None and (t.text or '').strip():
                                chosen_rPr = r.find(_qn('w:rPr'))
                                break
                        except Exception:
                            continue
                if chosen_rPr is None:
                    chosen_rPr = ref_para._element.find('.//' + _qn('w:rPr'))
                if chosen_rPr is not None:
                    rPr = _copy.deepcopy(chosen_rPr)
                    # Remove expanded/condensed character spacing
                    spacing_el = rPr.find(_qn('w:spacing'))
                    if spacing_el is not None:
                        rPr.remove(spacing_el)
            except Exception:
                pass

            if italic:
                try:
                    rPr.append(OxmlElement('w:i'))
                except Exception:
                    pass

            run_el.insert(0, rPr)
            t_el = OxmlElement('w:t')
            t_el.set(_qn('xml:space'), 'preserve')
            t_el.text = text
            run_el.append(t_el)
            new_p.append(run_el)
            ref_para._element.addnext(new_p)
            return new_p

        def _append_translation_linebreak(paragraph, text, italic=True):
            """Append translation as a new line within the same paragraph.

            This preserves original paragraph/table/cell structure better than inserting
            a brand new paragraph node (which can alter spacing/indent and break form layout).
            """
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn as _qn
            import copy as _copy

            txt = (text or '').strip()
            if not txt:
                return False

            try:
                # Add a soft line break within the same paragraph.
                paragraph.add_run('').add_break()

                tr = paragraph.add_run(txt)
                if italic:
                    try:
                        tr.italic = True
                    except Exception:
                        pass
                # Deep-copy the full rPr XML from the first non-empty source run
                # so ALL formatting (font, color, bold, underline, etc.) is preserved.
                # Strip character spacing (w:spacing) to avoid garbled wide letters.
                try:
                    src_rPr = None
                    for r in paragraph._element.findall('.//' + _qn('w:r')):
                        if r is tr._element:
                            continue
                        t_el = r.find(_qn('w:t'))
                        if t_el is not None and (t_el.text or '').strip():
                            rpr = r.find(_qn('w:rPr'))
                            if rpr is not None:
                                src_rPr = rpr
                                break
                    if src_rPr is not None:
                        new_rPr = _copy.deepcopy(src_rPr)
                        sp = new_rPr.find(_qn('w:spacing'))
                        if sp is not None:
                            new_rPr.remove(sp)
                        old_rPr = tr._element.find(_qn('w:rPr'))
                        if old_rPr is not None:
                            tr._element.remove(old_rPr)
                        tr._element.insert(0, new_rPr)
                except Exception:
                    pass
                return True
            except Exception:
                return False

        # Match dot/underscore/dash leaders (3+), ellipsis chars, and tab characters.
        # Lowered threshold to catch short fill-in-the-blank fields in forms.
        leader_re = re.compile(r"(\.{3,}|_{3,}|-{3,}|…+|\t+)")

        def _is_structural_text(text):
            """Return True if text is purely structural (leaders, tabs, punctuation)
            with no translatable word characters."""
            t = (text or "").strip()
            if not t:
                return True
            return not re.search(r'[\w\u00C0-\u1EF9]', t, flags=re.UNICODE)

        def _translate_preserve_form_leaders(text):
            """Translate text while preserving dot/underscore/dash leader runs
            and tab characters.

            This avoids breaking fill-in-the-blank lines in form-like DOCX templates.
            """
            raw = text or ""
            if not raw.strip():
                return raw
            # If text has no translatable word characters at all, return as-is
            if _is_structural_text(raw):
                return raw
            if not leader_re.search(raw):
                return _cleanup_translated_text(self._translate_with_retry(raw, target_lang))

            parts = leader_re.split(raw)
            out_parts = []
            for i, part in enumerate(parts):
                if i % 2 == 1:
                    # This is a leader/tab match — keep as-is
                    out_parts.append(part)
                    continue

                seg = part or ""
                if not seg.strip():
                    out_parts.append(seg)
                    continue
                # Skip translating segments that have no word characters.
                if _is_structural_text(seg):
                    out_parts.append(seg)
                    continue

                try:
                    out_parts.append(_cleanup_translated_text(self._translate_with_retry(seg, target_lang)))
                except ProviderRateLimitError:
                    raise
                except Exception:
                    if api_only:
                        raise
                    out_parts.append(seg)

            return _cleanup_translated_text("".join(out_parts))

        def _is_toc_paragraph(paragraph):
            try:
                style_name = str(getattr(getattr(paragraph, "style", None), "name", "") or "").lower()
                if "toc" in style_name:
                    return True
            except Exception:
                pass

            p_el = paragraph._element
            try:
                for node in p_el.xpath('.//*[local-name()="instrText"]'):
                    txt = "".join(node.itertext())
                    if "toc" in (txt or "").lower():
                        return True
            except Exception:
                pass

            try:
                for node in p_el.xpath('.//*[local-name()="fldSimple"]'):
                    for k, v in (node.attrib or {}).items():
                        if str(k).endswith("}instr") and "toc" in str(v or "").lower():
                            return True
            except Exception:
                pass

            return False

        def _flatten_hyperlinks_in_paragraph(paragraph):
            """Remove hyperlink wrappers but keep child runs in-place to preserve layout."""
            changed = False
            p_el = paragraph._element
            while True:
                links = list(p_el.xpath('./*[local-name()="hyperlink"]'))
                if not links:
                    break
                for link in links:
                    parent = link.getparent()
                    if parent is None:
                        continue
                    idx = parent.index(link)
                    for child in list(link):
                        parent.insert(idx, child)
                        idx += 1
                    parent.remove(link)
                    changed = True
            return changed

        def _normalize_toc_run_appearance(paragraph):
            # Strict layout mode: do not alter visual run styling.
            return

        def _normalize_toc_hyperlinks(document):
            touched = 0
            for para in iter_all_paragraphs(document):
                if not _is_toc_paragraph(para):
                    continue
                if _flatten_hyperlinks_in_paragraph(para):
                    touched += 1
                _normalize_toc_run_appearance(para)
            return touched

        def _normalize_generic_run_appearance(paragraph):
            for run in paragraph.runs:
                try:
                    run.underline = False
                except Exception:
                    pass
                try:
                    run.font.color.theme_color = None
                except Exception:
                    pass
                try:
                    run.font.color.rgb = RGBColor(0, 0, 0)
                except Exception:
                    pass

        def _strip_all_hyperlinks(document):
            touched = 0
            for para in iter_all_paragraphs(document):
                if _flatten_hyperlinks_in_paragraph(para):
                    touched += 1
                _normalize_generic_run_appearance(para)
            return touched

        def _paragraph_has_drawing(paragraph):
            try:
                return bool(paragraph._element.xpath('.//*[local-name()="drawing" or local-name()="pict"]'))
            except Exception:
                return False

        def _set_paragraph_text_preserve_runs(paragraph, new_text):
            runs = list(paragraph.runs)
            if not runs:
                paragraph.add_run(new_text or "")
                return

            # Never mutate drawing runs (logos/icons) directly.
            if _paragraph_has_drawing(paragraph):
                non_drawing_runs = []
                for r in runs:
                    try:
                        has_draw = bool(r._element.xpath('.//*[local-name()="drawing" or local-name()="pict"]'))
                    except Exception:
                        has_draw = False
                    if not has_draw:
                        non_drawing_runs.append(r)

                if not non_drawing_runs:
                    return

                target = None
                for r in non_drawing_runs:
                    if (r.text or "").strip():
                        target = r
                        break
                if target is None:
                    target = non_drawing_runs[0]

                target.text = new_text or ""
                for r in non_drawing_runs:
                    if r is not target:
                        r.text = ""
                return

            target = None
            for r in runs:
                if (r.text or "").strip():
                    target = r
                    break
            if target is None:
                target = runs[0]
            target.text = new_text or ""
            for r in runs:
                if r is not target:
                    r.text = ""

        def _is_in_table_cell(paragraph):
            try:
                parent = paragraph._element.getparent()
                return bool(parent is not None and (parent.tag or '').endswith('}tc'))
            except Exception:
                return False

        def _is_heading_paragraph(paragraph):
            try:
                style_name = str(getattr(getattr(paragraph, 'style', None), 'name', '') or '').lower()
                return style_name.startswith('heading')
            except Exception:
                return False

        def _normalize_heading_case(document):
            # Strict layout mode: do not rewrite heading content for casing.
            return 0

        def _normalize_table_header_text(text: str) -> str:
            t = (text or '').strip()
            norm = re.sub(r'\s+', ' ', t).lower()
            mapping = {
                'user id': 'User ID',
                'userid': 'User ID',
                'data type': 'Data Types',
                'data types': 'Data Types',
                'description': 'Description',
                'constraints': 'Constraints',
                'constraint': 'Constraints',
                'not nul': 'Not null',
            }
            return mapping.get(norm, t)

        def _normalize_table_layout_and_text(document):
            touched = 0
            term_map = {
                'school': 'Field',
                'mô tả': 'Description',
                'mo ta': 'Description',
                'ràng buộc': 'Constraints',
                'rang buoc': 'Constraints',
                'data types': 'Data Types',
                'data type': 'Data Types',
            }

            for table in document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            raw = ''.join((rr.text or '') for rr in para.runs)
                            if not raw.strip():
                                continue

                            fixed = _cleanup_translated_text(raw)
                            lowered = fixed.strip().lower()
                            if lowered in term_map:
                                fixed = term_map[lowered]

                            # Force remaining Vietnamese table text to English when target is English.
                            if str(target_lang).strip().lower().startswith('en') and re.search(r'[à-ỹđ]', fixed, flags=re.IGNORECASE):
                                try:
                                    fixed = _cleanup_translated_text(_translate_preserve_exact_lines(fixed))
                                except Exception:
                                    pass

                            if fixed != raw:
                                _set_paragraph_text_preserve_runs(para, fixed)
                                touched += 1
            return touched

        def _normalize_profile_tab_leaders(document):
            touched = 0
            key_re = re.compile(r'^(\s*)(student\s*id|email|class)\s*[:\-]?\s*(.*)$', flags=re.IGNORECASE)
            for para in iter_all_paragraphs(document):
                raw = ''.join((r.text or '') for r in para.runs)
                if not raw.strip():
                    continue
                # Keep mixed "Email ... Class ..." lines intact to avoid unwanted wrapping.
                low = raw.lower()
                if 'email' in low and 'class' in low:
                    continue
                if '\t' in raw and not re.search(r'\.{3,}', raw):
                    continue

                m = key_re.match(raw.strip())
                if not m:
                    continue

                label = m.group(2)
                rest = m.group(3) or ''
                value = re.sub(r'^[\.\-_:\s]+', '', rest).strip()
                if not value:
                    continue

                new_text = f"{label.title()}:\t{value}"
                if new_text != raw:
                    _set_paragraph_text_preserve_runs(para, new_text)
                    touched += 1

                try:
                    para.paragraph_format.tab_stops.add_tab_stop(Inches(5.6), WD_TAB_ALIGNMENT.LEFT, WD_TAB_LEADER.DOTS)
                except Exception:
                    pass
            return touched

        def _center_inline_images(document):
            centered = 0
            for para in iter_all_paragraphs(document):
                has_drawing = False
                try:
                    for run in para.runs:
                        dr = run._element.xpath('.//*[local-name()="drawing"]')
                        if dr:
                            has_drawing = True
                            break
                except Exception:
                    has_drawing = False
                if has_drawing:
                    try:
                        para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        centered += 1
                    except Exception:
                        pass
            return centered

        def _force_remaining_phrase_fixes(document):
            """Final safety pass for specific phrases that must be translated."""
            touched = 0
            replacements = [
                (r"\bKHOA\s+CÔNG\s+NGHỆ\s+THÔNG\s+TIN\b", "FACULTY OF INFORMATION TECHNOLOGY"),
                (r"\bKHOA\s+CONG\s+NGHE\s+THONG\s+TIN\b", "FACULTY OF INFORMATION TECHNOLOGY"),
                (r"Gửi\s+lại\s+phiếu\s+đăng\s+ký\s+qua\s+Email\s*:", "Resubmit the registration form via Email:"),
                (r"Gui\s+lai\s+phieu\s+dang\s+ky\s+qua\s+Email\s*:", "Resubmit the registration form via Email:"),
            ]

            for para in iter_all_paragraphs(document):
                if _paragraph_has_drawing(para):
                    continue
                raw = ''.join((r.text or '') for r in para.runs)
                if not raw.strip():
                    continue
                fixed = raw
                for pat, rep in replacements:
                    fixed = re.sub(pat, rep, fixed, flags=re.IGNORECASE)
                fixed = _cleanup_translated_text(fixed)
                if fixed != raw:
                    _set_paragraph_text_preserve_runs(para, fixed)
                    touched += 1
            return touched

        def translate_paragraph_runs(paragraph, idx=None, total=None):
            """Translate a paragraph, preserving per-run formatting.

            Uses format-group strategy: runs with different formatting
            (bold, italic, color …) are translated independently so each
            retains its original style.
            """
            runs = list(paragraph.runs)
            if not runs:
                return
            if _paragraph_has_drawing(paragraph):
                return

            original_texts = [r.text or "" for r in runs]
            paragraph_text = "".join(original_texts)
            if not paragraph_text.strip():
                return

            if bi_mode == 'newline':
                # Bilingual newline: keep original, translate full paragraph below
                # Skip purely structural paragraphs (only dots/leaders/tabs)
                if _is_structural_text(paragraph_text):
                    return
                try:
                    translated_para = _translate_preserve_form_leaders(paragraph_text)
                except ProviderRateLimitError:
                    raise
                except Exception as e:
                    print(f"Translator failed for paragraph: {e}")
                    if api_only:
                        raise
                    translated_para = paragraph_text
                if (translated_para or '').strip() and (translated_para or '').strip() != paragraph_text.strip():
                    # Check if inside a table cell — use linebreak to keep
                    # cell layout intact; otherwise use a new paragraph.
                    from docx.oxml.ns import qn as _qn2
                    parent_tag = ''
                    try:
                        parent_tag = paragraph._element.getparent().tag or ''
                    except Exception:
                        pass
                    if parent_tag.endswith('}tc'):
                        _append_translation_linebreak(paragraph, translated_para, italic=False)
                    else:
                        new_p = _insert_paragraph_after(paragraph, translated_para, italic=False)
                        try:
                            _seen_para_elems.add(id(new_p))
                        except Exception:
                            pass
            elif bi_mode == 'inline':
                try:
                    translated_para = _translate_preserve_form_leaders(paragraph_text)
                except ProviderRateLimitError:
                    raise
                except Exception as e:
                    print(f"Translator failed for paragraph: {e}")
                    if api_only:
                        raise
                    translated_para = paragraph_text
                t = (translated_para or '').strip()
                if t and t != paragraph_text.strip():
                    # In table cells, enforce strict bilingual separator requested by user: " | "
                    # to guarantee output style like: Username | Ten dang nhap
                    d = '|' if _is_in_table_cell(paragraph) else self._normalize_bilingual_delimiter(bilingual_delimiter)
                    last_run = None
                    for r in reversed(runs):
                        if (r.text or '').strip():
                            last_run = r
                            break
                    if last_run is None and runs:
                        last_run = runs[-1]
                    if last_run:
                        spacer = "" if (last_run.text or "").endswith((" ", "\t")) else " "
                        last_run.text = f"{last_run.text or ''}{spacer}{d} {t}"
            else:
                # Normal mode: use format-group translation for best formatting preservation
                _translate_format_groups(paragraph, _translate_preserve_form_leaders)

            if progress_callback and idx is not None and total is not None:
                progress_callback(10 + int((idx / total) * 70), f"Translating paragraph {idx+1}/{total}")

        # Avoid processing the same paragraph multiple times.
        # This can happen in tables with merged cells, where python-docx exposes the same
        # underlying paragraph through multiple cell references.
        _seen_para_elems = set()

        def _seen_or_mark(paragraph):
            try:
                key = id(paragraph._element)
            except Exception:
                key = id(paragraph)
            if key in _seen_para_elems:
                return True
            _seen_para_elems.add(key)
            return False

        # Preserve TOC placement while removing hyperlink wrappers only.
        try:
            _normalize_toc_hyperlinks(doc)
        except Exception:
            pass

        # Body paragraphs
        paragraphs = [p for p in doc.paragraphs]
        # Translate body paragraphs using the format-group approach.
        # For paragraphs with uniform formatting, translates the whole paragraph at once.
        # For paragraphs with mixed formatting (bold, color, etc.), translates each
        # format-group independently so each keeps its visual style.
        from concurrent.futures import as_completed

        total_work = 0
        completed = 0

        # Count workload
        body_paras = []
        for para in paragraphs:
            if _seen_or_mark(para):
                continue
            paragraph_text = "".join([r.text or "" for r in para.runs])
            if not paragraph_text.strip():
                continue
            body_paras.append(para)
        total_work = max(1, len(body_paras))

        for para in body_paras:
            try:
                if _paragraph_has_drawing(para):
                    continue
                original_texts = [r.text or "" for r in para.runs]
                paragraph_text = "".join(original_texts)

                if bi_mode == 'inline':
                    translated = _translate_preserve_exact_lines(paragraph_text)
                    t = (translated or '').strip()
                    if t and t != paragraph_text.strip():
                        d = self._normalize_bilingual_delimiter(bilingual_delimiter)
                        # Append delimiter + translation to the last non-empty run,
                        # preserving all original runs and their formatting.
                        last_run = None
                        for r in reversed(para.runs):
                            if (r.text or '').strip():
                                last_run = r
                                break
                        if last_run is None and para.runs:
                            last_run = para.runs[-1]
                        if last_run:
                            spacer = "" if (last_run.text or "").endswith((" ", "\t")) else " "
                            last_run.text = f"{last_run.text or ''}{spacer}{d} {t}"
                elif bi_mode == 'newline':
                    # Skip purely structural paragraphs (only dots/leaders/tabs)
                    if _is_structural_text(paragraph_text):
                        pass
                    else:
                        translated = _translate_preserve_exact_lines(paragraph_text)
                        if (translated or '').strip() and (translated or '').strip() != paragraph_text.strip():
                            # Check if inside a table cell
                            from docx.oxml.ns import qn as _qn3
                            parent_tag = ''
                            try:
                                parent_tag = para._element.getparent().tag or ''
                            except Exception:
                                pass
                            if parent_tag.endswith('}tc'):
                                _append_translation_linebreak(para, translated, italic=False)
                            else:
                                new_p = _insert_paragraph_after(para, translated, italic=False)
                                try:
                                    _seen_para_elems.add(id(new_p))
                                except Exception:
                                    pass
                else:
                    # Normal: use format-group translation for best formatting preservation
                    _translate_format_groups(para, _translate_preserve_exact_lines)
            except ProviderRateLimitError:
                print("Provider rate limit detected during paragraph processing, aborting job.")
                raise
            except Exception as e:
                print(f"Paragraph translation failed: {e}")
                if api_only:
                    raise
            completed += 1
            if progress_callback:
                progress_callback(
                    10 + int((completed / total_work) * 70),
                    f"Translating paragraph {completed}/{total_work}",
                )

        # Tables: translate cell-by-cell, preserve cell formatting
        for table in doc.tables:
            for r in range(len(table.rows)):
                for c in range(len(table.columns)):
                    cell = table.rows[r].cells[c]
                    for p_idx, p in enumerate(cell.paragraphs):
                        if _seen_or_mark(p):
                            continue
                        translate_paragraph_runs(p, p_idx, len(cell.paragraphs))

        # Headers and footers
        try:
            for section in doc.sections:
                header = section.header
                for p_idx, p in enumerate(header.paragraphs):
                    if _seen_or_mark(p):
                        continue
                    translate_paragraph_runs(p, p_idx, len(header.paragraphs))
                footer = section.footer
                for p_idx, p in enumerate(footer.paragraphs):
                    if _seen_or_mark(p):
                        continue
                    translate_paragraph_runs(p, p_idx, len(footer.paragraphs))
        except Exception:
            # ignore headers/footers issues
            pass

        # Optional: OCR embedded images in DOCX.
        # When mode includes 'image', we replace embedded image bytes.
        # When mode includes 'text', we export OCR+translated text to a sidecar .txt.
        if ocr_images and self.ocr_translate_overlay:
            if progress_callback:
                progress_callback(82, "OCR images in DOCX...")

            protected_image_partnames = _collect_header_footer_image_partnames(doc)

            paras_to_scan = iter_all_paragraphs(doc)
            total_paras = len(paras_to_scan) or 1
            images_found = 0
            ocr_attempted = 0
            ocr_success = 0
            ocr_disabled = False

            # Collect replacements: zip internal path -> bytes
            image_replacements = {}
            # Collect OCR text export entries
            ocr_export_entries = []
            # Collect entries for text insertion into DOCX paragraphs (used by mode='both')
            text_insert_entries = []  # list of (paragraph, translated_text)
            # Collect in-place replacements for mode='text': (paragraph, rid, translated_text)
            text_replace_entries = []

            def _is_probably_logo_or_nontext(ocr_text: str) -> bool:
                """Heuristic: avoid replacing small/non-text images (logos/icons)."""
                raw = (ocr_text or '').strip()
                if not raw:
                    return True
                words = re.findall(r'\w+', raw, flags=re.UNICODE)
                # Very short detections are usually logo marks, not document text blocks.
                if len(words) <= 2 and len(raw) < 24:
                    return True
                return False

            for idx, para in enumerate(paras_to_scan):
                if ocr_disabled:
                    break
                rids = paragraph_image_rids(para)
                if not rids:
                    continue
                for rid in rids:
                    img_part = rid_to_image_part(para, rid)
                    if not img_part:
                        continue
                    partname = str(getattr(img_part, 'partname', '') or '').lstrip('/')
                    if partname and partname in protected_image_partnames:
                        continue
                    try:
                        blob = getattr(img_part, 'blob', None)
                        if not blob:
                            continue

                        images_found += 1

                        ext = image_part_ext(img_part)
                        tmp_name = f"docx_img_{uuid.uuid4().hex}{ext}"
                        tmp_path = os.path.join(self.upload_folder, tmp_name)
                        with open(tmp_path, 'wb') as f:
                            f.write(blob)

                        ocr_attempted += 1
                        try:
                            # Render translated text back onto the image.
                            # Returns (ocr_text, translated_text, png_bytes, recommended_mode)
                            ocr_text, translated_text, png_bytes, ai_recommended_mode = self.ocr_translate_overlay(
                                tmp_path,
                                'auto',
                                target_lang,
                                ocr_langs,
                            )
                        finally:
                            try:
                                os.remove(tmp_path)
                            except Exception:
                                pass

                        if not ocr_text or not str(ocr_text).strip():
                            continue

                        # Do not touch likely logo/non-text images.
                        if _is_probably_logo_or_nontext(ocr_text):
                            continue

                        per_mode = mode
                        if mode == 'auto':
                            per_mode = _auto_pick_mode(ocr_text, translated_text, ai_recommended_mode)

                        print(f"  [IMAGE #{images_found}] OCR={len(ocr_text)}chars, AI_class={ai_recommended_mode}, per_mode={per_mode}")

                        # Export OCR text + translation for editing when requested
                        try:
                            if per_mode in ('text', 'both'):
                                partname = str(getattr(img_part, 'partname', '') or '').lstrip('/')
                                ocr_export_entries.append({
                                    'image': partname or '(embedded image)',
                                    'ocr_text': (ocr_text or '').strip(),
                                    'translated_text': (translated_text or '').strip(),
                                })
                                normalized_translated = _normalize_ocr_text_for_docx((translated_text or '').strip())

                                if per_mode == 'text' and mode == 'text':
                                    # EXPLICIT text mode: replace image with text in-place
                                    text_replace_entries.append((para, rid, normalized_translated))
                                else:
                                    # AUTO or 'both': keep original image + insert text below
                                    text_insert_entries.append((para, normalized_translated))
                        except Exception:
                            pass

                        if per_mode in ('image', 'both') and png_bytes and len(png_bytes) > 100:
                            # Replace the original embedded image bytes in the resulting docx.
                            try:
                                if partname:
                                    new_bytes = _overlay_bytes_to_original_format(png_bytes, ext)
                                    image_replacements[partname] = new_bytes
                                    ocr_success += 1
                            except Exception:
                                continue
                    except Exception as e:
                        # If Tesseract is missing/unavailable AND AI Vision is also failing,
                        # stop trying further images to avoid repeated failures.
                        msg = str(e).lower()
                        if ('tesseract' in msg and ('not installed' in msg or 'path' in msg)) or \
                           ('ocr unavailable' in msg):
                            ocr_disabled = True
                            if progress_callback:
                                progress_callback(85, "Skipping DOCX image OCR (OCR not available)")
                            break
                        if 'ai provider' in msg and ('not configured' in msg or 'rate' in msg):
                            ocr_disabled = True
                            if progress_callback:
                                progress_callback(85, f"Skipping DOCX image OCR: {e}")
                            break
                        # Otherwise, continue best-effort.
                        print(f"DOCX image OCR error (continuing): {e}")
                        continue

                if progress_callback and (idx % 10 == 0):
                    # Keep progress moving while OCR runs
                    progress_callback(82 + int((idx / total_paras) * 10), f"OCR scanning {idx+1}/{total_paras}")

            # Optional fallback package scan for floating/textbox/header images.
            # Disabled by default to avoid accidental logo/header replacement.
            try:
                pkg_scan_enabled = str(os.getenv('DOCX_OCR_PACKAGE_SCAN', '0')).strip().lower() in ('1', 'true', 'yes', 'on')
                if pkg_scan_enabled and (not ocr_disabled) and mode in ('image', 'both', 'auto'):
                    # Collect all image parts from the package
                    pkg = getattr(getattr(doc, 'part', None), 'package', None)
                    parts = list(getattr(pkg, 'parts', []) or [])

                    extra_attempted = 0
                    extra_replaced = 0
                    for part in parts:
                        try:
                            ct = str(getattr(part, 'content_type', '') or '').lower()
                            if not ct.startswith('image/'):
                                continue
                            partname = str(getattr(part, 'partname', '') or '').lstrip('/')
                            if not partname:
                                continue
                            if partname in protected_image_partnames:
                                continue
                            # Skip already processed via paragraph scan
                            if partname in image_replacements:
                                continue
                            blob = getattr(part, 'blob', None)
                            if not blob:
                                continue

                            ext = image_part_ext(part)
                            tmp_name = f"docx_img_pkg_{uuid.uuid4().hex}{ext}"
                            tmp_path = os.path.join(self.upload_folder, tmp_name)
                            with open(tmp_path, 'wb') as f:
                                f.write(blob)

                            extra_attempted += 1
                            try:
                                ocr_text, translated_text, png_bytes, ai_recommended_mode = self.ocr_translate_overlay(
                                    tmp_path,
                                    'auto',
                                    target_lang,
                                    ocr_langs,
                                )
                            finally:
                                try:
                                    os.remove(tmp_path)
                                except Exception:
                                    pass

                            if not ocr_text or not str(ocr_text).strip():
                                continue

                            if _is_probably_logo_or_nontext(ocr_text):
                                continue

                            per_mode = mode
                            if mode == 'auto':
                                per_mode = _auto_pick_mode(ocr_text, translated_text, ai_recommended_mode)

                            # For package-level scan we only apply overlay replacements.
                            if per_mode in ('image', 'both') and png_bytes and len(png_bytes) > 100:
                                try:
                                    new_bytes = _overlay_bytes_to_original_format(png_bytes, ext)
                                    image_replacements[partname] = new_bytes
                                    extra_replaced += 1
                                except Exception:
                                    continue
                        except Exception as e:
                            msg = str(e).lower()
                            if ('tesseract' in msg and ('not installed' in msg or 'path' in msg)) or \
                               ('ocr unavailable' in msg):
                                ocr_disabled = True
                                if progress_callback:
                                    progress_callback(85, "Skipping DOCX image OCR (OCR not available)")
                                break
                            if 'ai provider' in msg and ('not configured' in msg or 'rate' in msg):
                                ocr_disabled = True
                                if progress_callback:
                                    progress_callback(85, f"Skipping DOCX image OCR: {e}")
                                break
                            continue

                    if progress_callback and (extra_attempted or extra_replaced):
                        progress_callback(
                            92,
                            f"DOCX OCR (package scan): attempted={extra_attempted}, replaced={extra_replaced}",
                        )
            except Exception:
                pass

            if progress_callback:
                if images_found <= 0:
                    progress_callback(92, "DOCX OCR: no embedded images found")
                else:
                    progress_callback(
                        92,
                        f"DOCX OCR: found={images_found}, attempted={ocr_attempted}, replaced={ocr_success}",
                    )

            # Case 1: replace image by text at the same location
            try:
                if text_replace_entries:
                    for para, rid, trans_text in text_replace_entries:
                        if not trans_text or not trans_text.strip():
                            continue
                        try:
                            replace_image_with_text(para, rid, trans_text)
                        except Exception:
                            continue
            except Exception:
                pass

        # Insert translated text as paragraphs after image paragraphs in the DOCX
        try:
            if ocr_images and 'text_insert_entries' in locals() and text_insert_entries:
                from docx.oxml import OxmlElement
                from docx.oxml.ns import qn as _qn

                # Insert in REVERSE order so addnext doesn't shift positions
                for para, trans_text in reversed(text_insert_entries):
                    if not trans_text or not trans_text.strip():
                        continue
                    try:
                        # Split long text into multiple paragraphs for readability
                        text_paragraphs = [p.strip() for p in trans_text.split('\n') if p.strip()]
                        if not text_paragraphs:
                            text_paragraphs = [trans_text.strip()]

                        # Insert paragraphs in reverse so they appear in correct order after the image
                        for t_idx, t_para in enumerate(reversed(text_paragraphs)):
                            new_p = OxmlElement('w:p')

                            # Run with styled text
                            run = OxmlElement('w:r')

                            t_el = OxmlElement('w:t')
                            t_el.set(_qn('xml:space'), 'preserve')
                            t_el.text = t_para
                            run.append(t_el)
                            new_p.append(run)

                            para._element.addnext(new_p)
                    except Exception:
                        continue
        except Exception:
            pass

        # Targeted finishing pass requested by user (non-structural).
        # IMPORTANT: In bilingual modes, skip these post-processors because they can
        # overwrite already-generated "Original | Translation" text inside table cells.
        try:
            if bi_mode in ('inline', 'newline'):
                link_count = 0
                leader_count = 0
                table_count = 0
                img_count = 0
                forced_count = 0
                if progress_callback:
                    progress_callback(96, "DOCX targeted fixes skipped in bilingual mode")
            else:
                link_count = _strip_all_hyperlinks(doc)
                leader_count = _normalize_profile_tab_leaders(doc)
                table_count = _normalize_table_layout_and_text(doc)
                img_count = _center_inline_images(doc)
                forced_count = _force_remaining_phrase_fixes(doc)
                if progress_callback:
                    progress_callback(
                        96,
                        (
                            f"DOCX targeted fixes: links={link_count}, "
                            f"leaders={leader_count}, table={table_count}, images={img_count}, forced={forced_count}"
                        ),
                    )
        except Exception as e:
            if progress_callback:
                progress_callback(96, f"DOCX targeted fixes skipped: {e}")

        # Ensure output filename has .docx extension
        output_filename = f"translated_{os.path.basename(file_path)}"
        if not output_filename.lower().endswith('.docx'):
            output_filename += '.docx'
        output_path = os.path.join(self.download_folder, output_filename)

        # Save and validate
        doc.save(output_path)

        # NOTE: Sidecar OCR text export removed — translated text is now
        # inserted directly into the DOCX as paragraphs after each image.

        # If we rendered overlays, patch the embedded image bytes inside the saved DOCX (zip container)
        try:
            if ocr_images and mode in ('image', 'both', 'auto') and 'image_replacements' in locals() and image_replacements:
                if progress_callback:
                    progress_callback(93, "Applying translated overlays to DOCX images...")
                tmp_out = output_path + ".tmp"
                with zipfile.ZipFile(output_path, 'r') as zin, zipfile.ZipFile(tmp_out, 'w') as zout:
                    for item in zin.infolist():
                        data = zin.read(item.filename)
                        repl = image_replacements.get(item.filename)
                        if repl is not None:
                            data = repl
                        zout.writestr(item, data)
                # Replace original
                try:
                    os.replace(tmp_out, output_path)
                except Exception:
                    # Best-effort fallback
                    try:
                        os.remove(output_path)
                    except Exception:
                        pass
                    os.rename(tmp_out, output_path)
        except Exception as e:
            # If patching fails, keep the DOCX as saved (text translations still present)
            if progress_callback:
                progress_callback(94, f"DOCX image overlay patch failed: {e}")

        # Validate produced DOCX — if invalid, write a plain text fallback to avoid corrupt file being returned
        try:
            # Try opening the saved file with python-docx to validate
            docx.Document(output_path)
        except Exception as e:
            # Create a text fallback containing translated paragraphs and table text
            if progress_callback:
                progress_callback(95, "DOCX validation failed, writing fallback text file")
            fallback_filename = output_filename
            if not fallback_filename.lower().endswith('.txt'):
                fallback_filename += '.txt'
            fallback_path = os.path.join(self.download_folder, fallback_filename)
            # Collect text from doc
            lines = []
            for p in doc.paragraphs:
                lines.append(p.text)
            for t in doc.tables:
                for row in t.rows:
                    for cell in row.cells:
                        lines.append(cell.text)
            with open(fallback_path, 'w', encoding='utf-8') as f:
                f.write("NOTE: DOCX creation failed on server. Showing plain text fallback below.\n\n")
                f.write('\n'.join(lines))
            output_path = fallback_path

        if progress_callback:
            progress_callback(100, "Completed")
        return output_path

    def _process_xlsx(self, file_path, target_lang, progress_callback=None):
        # Translate in-place to preserve styles, merged cells, formulas, column widths, etc.
        wb = openpyxl.load_workbook(file_path)

        # Count total cells (rough progress)
        total_cells = 0
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            for _ in ws.iter_rows():
                total_cells += 1
        total_cells = total_cells or 1

        # Collect cells to translate
        to_translate = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            for row in ws.iter_rows():
                for cell in row:
                    try:
                        is_formula = (cell.data_type == 'f') or (
                            isinstance(cell.value, str) and cell.value.startswith("=")
                        )
                    except Exception:
                        is_formula = False

                    if (not is_formula) and isinstance(cell.value, str) and cell.value.strip():
                        to_translate.append(cell)

        total = len(to_translate) or 1
        processed = 0
        # Translate cells in parallel
        with self._executor_cls(max_workers=self.concurrency) as ex:
            futures = {ex.submit(self._translate_with_retry, cell.value, target_lang): cell for cell in to_translate}
            for fut in futures:
                try:
                    translated = fut.result()
                    cell = futures[fut]
                    cell.value = translated
                except ProviderRateLimitError:
                    print("Provider rate limit detected during cell translation, aborting job.")
                    raise
                except Exception as e:
                    print(f"Cell translation failed: {e}")
                processed += 1
                if progress_callback:
                    progress_callback(10 + int((processed / total) * 80), f"Translating cells {processed}/{total}")

        # Ensure output filename has .xlsx extension
        output_filename = f"translated_{os.path.basename(file_path)}"
        if not output_filename.lower().endswith('.xlsx'):
            output_filename += '.xlsx'
        output_path = os.path.join(self.download_folder, output_filename)
        wb.save(output_path)
        if progress_callback:
            progress_callback(100, "Completed")
        return output_path

    def _process_txt(self, file_path, target_lang, progress_callback=None):
        with open(file_path, 'r', encoding='utf-8') as f:
            text = f.read()
        if progress_callback:
            progress_callback(25, "Translating text file...")

        # Split into paragraphs then chunk long paragraphs to avoid provider length limits
        paras = [p.strip() for p in re.split(r'\n{2,}', text) if p.strip()]
        max_chars = 3000
        chunks = []
        for p in paras:
            if len(p) <= max_chars:
                chunks.append(p)
            else:
                parts = re.split(r'(?<=[.!?])\s+', p)
                cur = ''
                for part in parts:
                    if len(cur) + len(part) + 1 <= max_chars:
                        cur = (cur + ' ' + part).strip() if cur else part
                    else:
                        if cur:
                            chunks.append(cur)
                        cur = part
                if cur:
                    # If still too long, slice it
                    while len(cur) > max_chars:
                        chunks.append(cur[:max_chars])
                        cur = cur[max_chars:]
                    if cur:
                        chunks.append(cur)

        # Translate chunks in parallel
        translated_parts = []
        with self._executor_cls(max_workers=self.concurrency) as ex:
            futures = [ex.submit(self._translate_with_retry, c, target_lang) for c in chunks]
            total = len(futures) or 1
            for i, fut in enumerate(futures, start=1):
                try:
                    res = fut.result()
                    translated_parts.append(res)
                except ProviderRateLimitError as e:
                    try:
                        print("Provider rate limit detected during text file translation, aborting job.")
                    except UnicodeEncodeError:
                        print("Provider rate limit detected during text file translation, aborting job.")
                    raise
                except Exception as e:
                    print(f"Chunk translation failed: {e}")
                    translated_parts.append('')
                if progress_callback:
                    progress_callback(25 + int((i / total) * 70), f"Translating chunk {i}/{total}")

        translated_text = '\n\n'.join(translated_parts)

        output_filename = f"translated_{os.path.basename(file_path)}"
        if not output_filename.lower().endswith('.txt'):
            output_filename += '.txt'
        output_path = os.path.join(self.download_folder, output_filename)
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(translated_text)
        if progress_callback:
            progress_callback(100, "Completed")
        return output_path

    def _sanitize_text(self, text: str) -> str:
        if not isinstance(text, str):
            try:
                text = str(text)
            except Exception:
                return ''
        # Normalize unicode and remove control characters that break XML/docx
        text = unicodedata.normalize('NFC', text)
        # Remove C0 control characters except tab/newline/carriage return
        text = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F]', '', text)
        # Collapse weird zero-width/formatting if any
        text = re.sub(r'[\u200B-\u200F\u2028\u2029]', ' ', text)
        return text

    def _translate_text(self, text, target_lang):
        # Use injected translator with retry/backoff
        if self.translator:
            out = self._translate_with_retry(text, target_lang)
            return self._sanitize_text(out)
        raise RuntimeError("Translator is not configured")
