import os
import sys
import copy
from typing import Iterable

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def iter_all_paragraphs(doc: Document):
    # body
    for p in doc.paragraphs:
        yield p

    # tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p

    # headers/footers
    try:
        for section in doc.sections:
            for p in section.header.paragraphs:
                yield p
            for p in section.footer.paragraphs:
                yield p
    except Exception:
        return


def has_manual_break(paragraph) -> bool:
    br = qn('w:br')
    for r in getattr(paragraph, 'runs', []) or []:
        try:
            if r._element.findall('.//' + br):
                return True
        except Exception:
            pass
    return False


def split_bilingual_breaks(doc: Document) -> int:
    """Convert manual line-break bilingual paragraphs into two paragraphs.

    Expected structure produced by older newline-mode:
      [original runs...] [run containing w:br] [translation runs...]

    We remove the break+translation from the original paragraph (so its layout is restored)
    and insert a new paragraph after it containing the translation text.
    """

    br = qn('w:br')
    fixed = 0
    seen = set()

    for p in list(iter_all_paragraphs(doc)):
        try:
            peid = id(p._element)
        except Exception:
            peid = id(p)
        if peid in seen:
            continue
        seen.add(peid)

        if not has_manual_break(p):
            continue

        runs = list(getattr(p, 'runs', []) or [])
        if not runs:
            continue

        break_idx = None
        for i, r in enumerate(runs):
            try:
                if r._element.findall('.//' + br):
                    break_idx = i
                    break
            except Exception:
                continue

        if break_idx is None:
            continue

        # Translation text = all text after the break run
        trans_text = ''.join((r.text or '') for r in runs[break_idx + 1 :]).strip()
        if not trans_text:
            continue

        # Capture formatting from the source paragraph/run (before we remove anything)
        src_pPr = None
        src_rPr = None
        try:
            src_pPr = p._element.find(qn('w:pPr'))
        except Exception:
            src_pPr = None
        try:
            # Prefer a run with explicit color formatting
            for r in runs[:break_idx]:
                if not (r.text or '').strip():
                    continue
                rpr = r._element.find(qn('w:rPr'))
                if rpr is not None and rpr.find(qn('w:color')) is not None:
                    src_rPr = rpr
                    break
            # Fallback: first non-empty run
            if src_rPr is None:
                for r in runs[:break_idx]:
                    if (r.text or '').strip():
                        src_rPr = r._element.find(qn('w:rPr'))
                        break
        except Exception:
            src_rPr = None

        # Heuristic: break is expected at the end (translation appended).
        # If there is substantial original-looking text after break, skip.
        # (We treat any content after break as translation; safer to require that break is near end.)
        if break_idx < max(0, len(runs) - 4):
            # Still proceed, but only if content after break contains latin letters (likely translation)
            if not any('a' <= ch.lower() <= 'z' for ch in trans_text):
                continue

        # Remove break run and all runs after it from the original paragraph
        try:
            for r in runs[break_idx:]:
                try:
                    r_el = r._element
                    parent = r_el.getparent()
                    if parent is not None:
                        parent.remove(r_el)
                except Exception:
                    try:
                        r.text = ''
                    except Exception:
                        pass
        except Exception:
            continue

        # Insert a new paragraph right after with the translation text
        try:
            new_p = OxmlElement('w:p')

            # Copy pPr to preserve layout, but remove numbering to avoid duplicating bullets.
            if src_pPr is not None:
                ppr_copy = copy.deepcopy(src_pPr)
                try:
                    numPr = ppr_copy.find(qn('w:numPr'))
                    if numPr is not None:
                        ppr_copy.remove(numPr)
                except Exception:
                    pass
                new_p.append(ppr_copy)

            run_el = OxmlElement('w:r')
            if src_rPr is not None:
                try:
                    run_el.append(copy.deepcopy(src_rPr))
                except Exception:
                    pass
            t_el = OxmlElement('w:t')
            t_el.set(qn('xml:space'), 'preserve')
            t_el.text = trans_text
            run_el.append(t_el)
            new_p.append(run_el)

            p._element.addnext(new_p)
            fixed += 1
            seen.add(id(new_p))
        except Exception:
            continue

    return fixed


def main(argv: list[str]) -> int:
    src = argv[1] if len(argv) > 1 else r"backend/downloads/translated_testfile_-_Copy (10).docx"
    out = argv[2] if len(argv) > 2 else None

    if not out:
        base, ext = os.path.splitext(src)
        out = base + "_fixed" + (ext or '.docx')

    doc = Document(src)
    fixed = split_bilingual_breaks(doc)
    doc.save(out)

    print(f"input:  {src}")
    print(f"output: {out}")
    print(f"fixed_paragraphs: {fixed}")
    return 0


if __name__ == '__main__':
    raise SystemExit(main(sys.argv))
