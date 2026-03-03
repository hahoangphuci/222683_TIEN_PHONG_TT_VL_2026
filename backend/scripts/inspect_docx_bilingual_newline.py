import sys
from docx import Document
from docx.oxml.ns import qn


def para_has_break(p) -> bool:
    br_tag = qn('w:br')
    for r in getattr(p, 'runs', []) or []:
        try:
            if r._element.findall('.//' + br_tag):
                return True
        except Exception:
            pass
    return False


def scan_docx(path: str, limit: int = 120) -> int:
    doc = Document(path)
    suspicious = []

    # Body
    for i, p in enumerate(doc.paragraphs):
        txt = (p.text or '')
        if '→' in txt or '\n' in txt or para_has_break(p):
            suspicious.append(("body", str(i), txt))

    # Tables
    for ti, t in enumerate(doc.tables):
        for ri, row in enumerate(t.rows):
            for ci, cell in enumerate(row.cells):
                for pi, p in enumerate(cell.paragraphs):
                    txt = (p.text or '')
                    if '→' in txt or '\n' in txt or para_has_break(p):
                        suspicious.append((f"table{ti}", f"{ri}:{ci}:{pi}", txt))

    print(f"file: {path}")
    print(f"body_paragraphs: {len(doc.paragraphs)}")
    print(f"suspicious_paragraphs: {len(suspicious)}")

    for kind, idx, txt in suspicious[:limit]:
        txt2 = txt.replace('\n', '\\n')
        if len(txt2) > 240:
            txt2 = txt2[:240] + '…'
        print(f"[{kind} {idx}] {txt2}")

    return 0


if __name__ == '__main__':
    p = sys.argv[1] if len(sys.argv) > 1 else r"backend/downloads/translated_testfile_-_Copy (10).docx"
    raise SystemExit(scan_docx(p))
